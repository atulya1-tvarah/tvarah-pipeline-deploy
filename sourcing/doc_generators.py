"""Generate sample 100/100 resume and analysis guide as Word (.docx) and PDF."""
from __future__ import annotations

import io
from typing import Any


# ---------------------------------------------------------------------------
# WORD — Perfect Resume
# ---------------------------------------------------------------------------

def build_perfect_resume_docx() -> bytes:
    """Return bytes of a .docx 100/100 sample resume."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.oxml

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Default paragraph font ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Color constants ──
    PRIMARY   = RGBColor(0x35, 0x33, 0x95)   # deep indigo
    MUTED     = RGBColor(0x62, 0x74, 0x8E)
    BLACK     = RGBColor(0x26, 0x26, 0x26)
    GREEN     = RGBColor(0x16, 0xA3, 0x4A)
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

    def _add_h1(text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = PRIMARY
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    def _add_contact_line(text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    def _add_section_heading(text: str) -> None:
        doc.add_paragraph()  # small gap
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = PRIMARY
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(4)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "353395")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_job(title: str, company: str, period: str, location: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(1)
        r_title = p.add_run(title)
        r_title.bold = True
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = BLACK
        r_sep = p.add_run("  •  ")
        r_sep.font.color.rgb = MUTED
        r_sep.font.size = Pt(9.5)
        r_co = p.add_run(company)
        r_co.font.size = Pt(10)
        r_co.font.color.rgb = PRIMARY
        r_co.bold = True
        # right-align period using tab
        p.add_run("\t")
        r_period = p.add_run(f"{period}  |  {location}")
        r_period.font.size = Pt(9)
        r_period.font.color.rgb = MUTED
        # Tab stop at right margin
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9072")   # ~16 cm from left
        tabs.append(tab)
        pPr.append(tabs)

    def _bullet(text: str, bold_prefix: str = "") -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent   = Inches(0.2)
        p.paragraph_format.space_after   = Pt(1)
        p.paragraph_format.space_before  = Pt(0)
        if bold_prefix:
            r = p.add_run(bold_prefix + " ")
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = GREEN
        r2 = p.add_run(text)
        r2.font.size = Pt(9.5)

    def _skill_line(label: str, detail: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        rb = p.add_run(label + ": ")
        rb.bold = True
        rb.font.size = Pt(9.5)
        rb.font.color.rgb = PRIMARY
        rd = p.add_run(detail)
        rd.font.size = Pt(9.5)

    def _edu_row(degree: str, inst: str, year: str, detail: str = "") -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)
        rb = p.add_run(degree + " — ")
        rb.bold = True
        rb.font.size = Pt(10)
        ri = p.add_run(inst)
        ri.font.size = Pt(10)
        ri.font.color.rgb = PRIMARY
        ri.bold = True
        p.add_run(f"  ({year})")
        if detail:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(0)
            rd = p2.add_run(detail)
            rd.font.size  = Pt(9)
            rd.font.color.rgb = MUTED

    # ═══════════════════════════════════════════════════
    # HEADER
    # ═══════════════════════════════════════════════════
    _add_h1("Arjun Mehta")
    _add_contact_line("Bangalore, India  •  arjun.mehta@gmail.com  •  +91 98765 43210")
    _add_contact_line("linkedin.com/in/arjunmehta  •  github.com/arjunmehta  •  arjunmehta.dev")

    # ═══════════════════════════════════════════════════
    # PROFESSIONAL SUMMARY
    # ═══════════════════════════════════════════════════
    _add_section_heading("Professional Summary")
    p = doc.add_paragraph()
    r = p.add_run(
        "Staff Engineer with 11+ years building high-scale distributed systems across e-commerce, fintech, and adtech. "
        "Led cross-functional teams of 8–12 engineers at FAANG-tier companies. "
        "Architected systems handling 2M+ QPS with sub-10ms p99 latency. "
        "2× patent holder, IEEE published author, CKA + AWS SAP certified."
    )
    r.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(2)

    # ═══════════════════════════════════════════════════
    # EXPERIENCE
    # ═══════════════════════════════════════════════════
    _add_section_heading("Experience")

    # --- Job 1 ---
    _add_job("Staff Engineer", "Google India", "Jan 2021 – Present  (3.5 yr)", "Bangalore")
    _bullet(
        "Architected a real-time ad-serving pipeline (Go, Kubernetes, Bigtable, Pub/Sub) handling 2M QPS "
        "with p99 latency < 8ms — reduced infrastructure cost by $1.4M/year.",
        "▶"
    )
    _bullet("Led a team of 8 engineers; mentored 3 SDEs promoted within 18 months.", "▶")
    _bullet(
        "Collaborated with Google NYC team on cross-region replication for 6 months (onsite); "
        "presented reliability roadmap to VP Engineering.",
        "▶"
    )
    _bullet("Won Google 'Peer Bonus' award (2022) and 'Spot Award' (2023).", "▶")
    p = doc.add_paragraph()
    r = p.add_run("Stack: Go · Kubernetes · Bigtable · Terraform · Prometheus · Pub/Sub · Spanner")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(1)

    # --- Job 2 ---
    _add_job("Senior Engineer → Staff Engineer", "Swiggy", "Mar 2018 – Dec 2020  (2.8 yr)", "Bangalore")
    _bullet(
        "Owned end-to-end delivery of a dynamic surge-pricing ML system (Python, XGBoost, Airflow, Kafka) "
        "reducing delivery cost by ₹3.2 Cr/month — sole architect, 9-month delivery.",
        "▶"
    )
    _bullet(
        "Built real-time order-matching microservices (FastAPI, Redis, PostgreSQL) serving 400K orders/day; "
        "reduced checkout latency by 28%.",
        "▶"
    )
    _bullet("Promoted from Senior to Staff Engineer in 14 months — fastest in the cohort.", "▶")
    _bullet("Won 'Engineer of the Year 2020' (company-wide award).", "▶")
    p = doc.add_paragraph()
    r = p.add_run("Stack: Python · FastAPI · XGBoost · Airflow · Kafka · Redis · PostgreSQL · Docker")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(1)

    # --- Job 3 ---
    _add_job("Senior Software Engineer", "Flipkart", "Jun 2015 – Feb 2018  (2.8 yr)", "Bangalore")
    _bullet(
        "Built the product search ranking pipeline (Elasticsearch, Python, Spark) "
        "improving click-through rate by 18% on 200M SKUs.",
        "▶"
    )
    _bullet("Conducted 120+ technical interviews; shortlisted 14 hires.", "▶")
    p = doc.add_paragraph()
    r = p.add_run("Stack: Python · Elasticsearch · Apache Spark · Hive · MySQL · Cassandra")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(1)

    # --- Job 4 ---
    _add_job("Software Engineer", "Infosys", "Jul 2013 – May 2015  (1.9 yr)", "Pune")
    _bullet("Developed REST APIs for a banking client (Java, Spring Boot, Oracle); delivered 0 P1 incidents over 18 months.", "▶")
    p = doc.add_paragraph()
    r = p.add_run("Stack: Java · Spring Boot · Oracle · Maven · Jenkins")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED

    # ═══════════════════════════════════════════════════
    # SKILLS
    # ═══════════════════════════════════════════════════
    _add_section_heading("Technical Skills")
    _skill_line("Languages",    "Python (9 yr, Expert)  ·  Go (4 yr, Advanced)  ·  Java (3 yr, Applied)  ·  SQL (10 yr, Expert)")
    _skill_line("Frameworks",   "FastAPI (4 yr, Advanced)  ·  Spring Boot (3 yr, Hands-On)  ·  Spark (4 yr, Advanced)")
    _skill_line("ML / Data",    "XGBoost (5 yr, Advanced)  ·  Scikit-learn (4 yr, Hands-On)  ·  Airflow (4 yr, Advanced)  ·  Kafka (5 yr, Advanced)")
    _skill_line("Infra / Cloud","Kubernetes (5 yr, Advanced)  ·  Terraform (3 yr, Advanced)  ·  AWS (4 yr, Advanced)  ·  GCP (3 yr, Hands-On)")
    _skill_line("Databases",    "PostgreSQL (8 yr, Expert)  ·  Redis (6 yr, Advanced)  ·  Bigtable (3 yr, Advanced)  ·  Cassandra (3 yr, Hands-On)")
    _skill_line("Coding",       "GitHub: github.com/arjunmehta  ·  LeetCode: 350+ problems solved  ·  Kaggle Expert rank")

    # ═══════════════════════════════════════════════════
    # CERTIFICATIONS
    # ═══════════════════════════════════════════════════
    _add_section_heading("Certifications")
    certs = [
        ("AWS Solutions Architect Professional", "Amazon Web Services", "2023"),
        ("Certified Kubernetes Administrator (CKA)", "CNCF", "2022"),
        ("Google Professional Machine Learning Engineer", "Google Cloud", "2021"),
        ("Databricks Certified Associate Developer for Apache Spark", "Databricks", "2020"),
    ]
    for name, issuer, year in certs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(0)
        rb = p.add_run(f"• {name}")
        rb.bold = True; rb.font.size = Pt(9.5)
        ri = p.add_run(f"  —  {issuer}  ({year})")
        ri.font.size = Pt(9); ri.font.color.rgb = MUTED

    # ═══════════════════════════════════════════════════
    # EDUCATION
    # ═══════════════════════════════════════════════════
    _add_section_heading("Education")
    _edu_row("M.Tech Computer Science", "IIT Bombay", "2011–2013",
             "CPI: 9.1 / 10  ·  Thesis: 'Distributed consensus optimisation in heterogeneous networks'")
    _edu_row("B.Tech Computer Engineering", "IIT Bombay", "2007–2011",
             "CPI: 8.7 / 10  ·  Awarded Institute Silver Medal for academic excellence")

    # ═══════════════════════════════════════════════════
    # PUBLICATIONS & PATENTS
    # ═══════════════════════════════════════════════════
    _add_section_heading("Publications & Patents")
    pubs = [
        "Mehta A. et al. — 'Adaptive Latency-Aware Scheduling in Real-Time Bidding Systems', IEEE INFOCOM 2023",
        "Mehta A., Kumar R. — 'Low-Latency Feature Store for Online ML Inference', KDD Workshop 2021",
        "Patent: 'Method and System for Dynamic Surge Pricing in Last-Mile Logistics' (IN Patent 412763, Filed 2020, Granted 2022)",
    ]
    for pub in pubs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(f"• {pub}")
        r.font.size = Pt(9.5)

    # ═══════════════════════════════════════════════════
    # EXTRA-CURRICULARS
    # ═══════════════════════════════════════════════════
    _add_section_heading("Achievements & Activities")
    extras = [
        "Smart India Hackathon 2012 — National Winner (Team of 4)",
        "Open-source contributor: 3 repos with 800+ GitHub stars; maintainer of 'pystream-ml' library",
        "Executive PG in AI & ML — upGrad / IIIT Bangalore (2020, Distinction)",
        "NSS Volunteer (2008–2011) — taught coding to 200+ school students",
        "Chess club captain, IIT Bombay (2009–2010)",
    ]
    for e in extras:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(f"• {e}")
        r.font.size = Pt(9.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# WORD — Analysis Guide
# ---------------------------------------------------------------------------

def build_analysis_guide_docx() -> bytes:  # noqa: C901
    """Return bytes of the Tvarah scoring guide .docx (full signal tables)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    PRIMARY  = RGBColor(0x35, 0x33, 0x95)
    MUTED    = RGBColor(0x62, 0x74, 0x8E)
    GREEN    = RGBColor(0x16, 0xA3, 0x4A)
    AMBER    = RGBColor(0xD9, 0x77, 0x06)
    RED      = RGBColor(0xDC, 0x26, 0x26)
    BLUE     = RGBColor(0x25, 0x63, 0xEB)
    HDR_BG   = "353395"   # hex for table header shading
    ALT_BG   = "F0F0FA"   # light lavender for alternate rows

    # ── Shared helpers ──────────────────────────────────────────────────────

    def _h1(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(20); r.font.color.rgb = PRIMARY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)

    def _h2(text: str) -> None:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PRIMARY
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "353395")
        pBdr.append(bot); pPr.append(pBdr)

    def _h3(text: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = PRIMARY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)

    def _body(text: str, indent: bool = False) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(2)
        if indent:
            p.paragraph_format.left_indent = Inches(0.2)

    def _bullet(text: str, color: RGBColor = MUTED) -> None:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.size = Pt(9); r.font.color.rgb = color
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)

    def _shade_cell(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _signal_table(headers: list[str], rows: list[list[str]],
                      col_widths: list[float] | None = None) -> None:
        """Insert a compact lookup table. Header row = dark indigo, alt rows shaded."""
        n_cols = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.style = "Table Grid"
        tbl.autofit = False
        # Column widths
        if col_widths:
            for i, w in enumerate(col_widths):
                for cell in tbl.columns[i].cells:
                    cell.width = Inches(w)
        # Header row
        hdr = tbl.rows[0]
        for i, h in enumerate(headers):
            c = hdr.cells[i]
            _shade_cell(c, HDR_BG)
            p = c.paragraphs[0]
            p.clear()
            run = p.add_run(h)
            run.bold = True; run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Data rows
        for ri, row_data in enumerate(rows):
            row = tbl.rows[ri + 1]
            if ri % 2 == 1:
                for ci in range(n_cols):
                    _shade_cell(row.cells[ci], ALT_BG)
            for ci, val in enumerate(row_data):
                c = row.cells[ci]
                p = c.paragraphs[0]
                p.clear()
                run = p.add_run(str(val))
                run.font.size = Pt(8.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _param_header(param: str, pts: str, stage: str,
                      stage_color: RGBColor | None = None) -> None:
        """Bold parameter name + pts badge + stage tag."""
        stage_color = stage_color or (
            AMBER if "recruiter" in stage.lower()
            else RED if "panel" in stage.lower()
            else GREEN
        )
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        rb = p.add_run(f"{param}  ")
        rb.bold = True; rb.font.size = Pt(10.5)
        rp = p.add_run(f"[{pts}]")
        rp.bold = True; rp.font.color.rgb = stage_color; rp.font.size = Pt(10.5)
        rs = p.add_run(f"  ·  {stage}")
        rs.font.size = Pt(9); rs.font.color.rgb = MUTED; rs.italic = True

    # ══════════════════════════════════════════
    # COVER
    # ══════════════════════════════════════════
    _h1("Tvarah Resume Analysis")
    _h1("Scoring Guide — v5")
    p = doc.add_paragraph()
    r = p.add_run(
        "Complete signal reference  ·  Experience 40 + Skills 45 + Education 15 = 100 pts  "
        "·  3-stage model: Resume Auto → Recruiter → Panel"
    )
    r.font.size = Pt(10); r.font.color.rgb = MUTED; r.italic = True

    # ══════════════════════════════════════════
    # ARCHITECTURE
    # ══════════════════════════════════════════
    _h2("Scoring Architecture")
    _body(
        "Three sections, three stages. Each section has a hard cap. Recruiter and panel params "
        "start at 0 on the resume stage — they only add points when recruiter/panel fill them in. "
        "A perfect resume alone can score ~99.5 pts; in practice most candidates need all 3 stages."
    )
    _signal_table(
        ["Section", "Resume Auto (max)", "Recruiter + (max)", "Panel + (max)", "Section Cap"],
        [
            ["Experience",  "40 pts  (8 params)",   "7 pts  (3 params)",    "—",                 "40 pts"],
            ["Skills",      "45 pts  (5 params)",   "7 pts  (2 params)",    "13 pts  (3 params)", "45 pts"],
            ["Education",   "14.5 pts  (7 params)", "1 pt   (1 param)",     "—",                 "15 pts"],
            ["TOTAL",       "99.5 pts raw",         "15 pts max",           "13 pts max",         "100 pts"],
        ],
        col_widths=[1.2, 1.6, 1.6, 1.5, 1.0],
    )

    # ══════════════════════════════════════════
    # EXPERIENCE
    # ══════════════════════════════════════════
    _h2("Experience Section — 40 pts")

    _param_header("Overall Experience", "4 pts", "Resume Auto")
    _body("Signal: total_experience_years — sum of all non-overlapping work durations.")
    _signal_table(
        ["Total YoE", "Score"],
        [
            ["10+ years",      "4.0"],
            ["6–10 years",     "3.0"],
            ["4–6 years",      "2.5"],
            ["2–4 years",      "2.0"],
            ["1–2 years",      "1.5"],
            ["< 1 year",       "1.0"],
            ["0 / not found",  "0.0"],
        ],
        col_widths=[2.5, 1.0],
    )

    _param_header("Career Breaks", "2 pts", "Resume Auto")
    _body("Signal: gaps > 3 months between consecutive roles. Counts number of breaks and longest gap.")
    _signal_table(
        ["Break Pattern", "Score"],
        [
            ["0 breaks",                 "2.0"],
            ["1 break  ≤ 6 months",      "1.5"],
            ["1 break  ≤ 12 months",     "1.0"],
            ["1 break  > 12 months",     "0.5"],
            ["2 breaks (any)",           "0.5"],
            ["3+ breaks",               "0.0  (+REJECT flag)"],
        ],
        col_widths=[2.5, 1.5],
    )

    _param_header("Career Progression", "4 pts", "Resume Auto + BERT")
    _body("Heuristic (title hierarchy) blended with BERT label. Blend ratio driven by confidence band.")
    _signal_table(
        ["BERT Label",  "Meaning",                               "Score"],
        [
            ["FAST_TRACK", "2+ levels up in < 5 years",              "4.0"],
            ["GROWING",    "Steady upward advance, at least 1 promo","3.0"],
            ["LATERAL",    "Same level, domain switch",               "2.0"],
            ["DECLINING",  "Downward or stagnant titles",             "0.5"],
        ],
        col_widths=[1.3, 2.8, 0.7],
    )
    _signal_table(
        ["Confidence Band",        "Blend"],
        [
            [">=0.65  (high)",      "60% BERT + 40% heuristic"],
            ["0.45-0.64  (medium)", "40% BERT + 60% heuristic"],
            ["< 0.45  (low)",       "20% BERT + 80% heuristic"],
        ],
        col_widths=[1.8, 2.5],
    )

    _param_header("Stability", "4 pts", "Resume Auto")
    _body("Signals: avg_tenure, hop_rate (jobs/year), short_stints count, loyalty_label. "
          "stability_score/5 x 4 = final pts.")
    _signal_table(
        ["Condition",                                     "Effect"],
        [
            ["short_stints >= 2  (each < 18 months)",     "Penalty flag"],
            ["hop_rate >= 1.5  (>= 1.5 jobs/year)",       "Penalty flag"],
            ["loyalty = LOW  (avg tenure < 18 months)",  "Penalty flag"],
            ["No flags, HIGH loyalty",                   "Full 4 pts"],
        ],
        col_widths=[3.0, 1.2],
    )

    _param_header("Company Tier", "6 pts", "Resume Auto")
    _body("Signal: best company across all work entries matched against COMPANY_DICTIONARY.")
    _signal_table(
        ["Tier",                     "Examples",                                                        "Pts"],
        [
            ["TIER_1  (FAANG/Unicorn)", "Google, Meta, Amazon, Microsoft, Apple, Flipkart, Swiggy, Zomato, CRED, Razorpay", "6"],
            ["TIER_2  (Strong MNC/product)", "Adobe, Oracle, SAP, Salesforce, Paytm, Ola, Nykaa, Urban Company",           "5"],
            ["TIER_3  (Mid-size product)", "Funded scale-ups, growing B2B/B2C startups",                                   "4"],
            ["TIER_4  (IT Services)", "TCS, Infosys, Wipro, HCL, Cognizant, Capgemini, Accenture",                        "3"],
            ["TIER_5  (Unknown/tiny)", "No recognisable brand or < 50 employees",                                         "1"],
        ],
        col_widths=[1.8, 2.8, 0.4],
    )

    _param_header("Awards & Recognition", "4 pts", "Resume Auto")
    _body("Signal: count of items in achievements[] — phrases like 'promoted', 'hackathon winner', "
          "'performance award', 'employee of the month'.")
    _signal_table(
        ["Count",  "Score"],
        [["0","0"],["1","1"],["2","2"],["3+","4"]],
        col_widths=[2.0, 1.0],
    )

    _h3("Project Scoring  (16 pts combined)")
    _body("Deterministic signal checks + LLM complexity rating (0-5). "
          "Final = 40% deterministic + 60% LLM for both P1 and P2.")

    _param_header("Project 1", "10 pts", "Resume Auto + LLM")
    _signal_table(
        ["Signal",             "Check",                                                   "Pts"],
        [
            ["project_type",       "Category tagged  (ML, Backend, Platform...)",               "1"],
            ["your_role",          "Role stated  (Lead Engineer, Owner, Architect...)",          "1"],
            ["problem_description","Description > 50 chars",                                    "1"],
            ["duration",           "Duration >= 3 months explicitly stated",                    "1"],
            ["skills_used",        ">= 1 skill/technology listed",                              "1"],
            ["domain",             "Business domain tagged  (FinTech, E-Commerce...)",          "1"],
            ["ownership_verb",     "Verb: Architected / Owned / Led / Designed",                "+1"],
            ["quantified_impact",  "Number + outcome  (%, Rs/$, ms latency, user count)",       "+1"],
            ["cross_functional",   "Cross-team / stakeholder coordination mentioned",           "+0.5"],
            ["LLM complexity",     "0-5 rating x 60% weight",                                  "0-6"],
        ],
        col_widths=[1.5, 2.8, 0.7],
    )

    _param_header("Project 2", "6 pts", "Resume Auto + LLM")
    _signal_table(
        ["Signal",             "Check"],
        [
            ["project_type",       "Category tagged"],
            ["your_role",          "Role stated"],
            ["problem_description","Description > 20 chars"],
            ["duration",           "Duration >= 3 months"],
            ["skills_used",        ">= 1 skill listed"],
            ["domain",             "Domain tagged"],
        ],
        col_widths=[1.5, 3.5],
    )

    _h3("Recruiter-Stage Parameters  (+7 pts)")

    _param_header("International Exposure", "2 pts", "Recruiter Stage")
    _body("Recruiter validates: onsite trips, global team mentions, multi-country client exposure, "
          "relocation history. Assigns 0-2 after phone screen.")

    _param_header("Stakeholder Management", "2 pts", "Recruiter Stage + BERT")
    _body("BERT prior  (NONE / INTERNAL / CLIENT_FACING / C_LEVEL)  blended with recruiter confirmation.")
    _signal_table(
        ["BERT Label",     "Evidence Pattern",                        "Score"],
        [
            ["C_LEVEL",        "Board/exec presentations, C-suite interaction", "2.0"],
            ["CLIENT_FACING",  "Direct client comms, vendor management",        "1.5"],
            ["INTERNAL",       "Cross-functional coordination only",            "1.0"],
            ["NONE",           "No stakeholder signals found",                  "0.0"],
        ],
        col_widths=[1.3, 2.8, 0.7],
    )

    _param_header("Mentorship Signal", "3 pts", "Recruiter Stage + BERT")
    _body("BERT prior  (NONE / IMPLIED / FORMAL / LEAD)  blended with recruiter confirmation.")
    _signal_table(
        ["BERT Label", "Evidence Pattern",                               "Score"],
        [
            ["LEAD",     "'Led team of N', delivery owner, hiring decisions", "3.0"],
            ["FORMAL",   "'Mentored X engineers', structured programme",      "2.0"],
            ["IMPLIED",  "'Guided juniors', informal coaching mentioned",      "1.0"],
            ["NONE",     "No mentorship signals",                             "0.0"],
        ],
        col_widths=[1.0, 3.2, 0.7],
    )

    # ══════════════════════════════════════════
    # SKILLS
    # ══════════════════════════════════════════
    _h2("Skills Section — 45 pts")

    _param_header("Skill List x Years", "20 pts", "Resume Auto")
    _body("Per skill: if evidence_level in {APPLIED, DEEP, EXPERT} add years_of_usage (capped at 10). "
          "Total / 30 x 20 = final score.")
    _signal_table(
        ["Evidence Level",  "Counted?",        "Cap per Skill"],
        [
            ["EXPERT",       "Yes — full",      "10 yrs"],
            ["DEEP",         "Yes — full",      "10 yrs"],
            ["APPLIED",      "Yes — full",      "10 yrs"],
            ["FOUNDATIONAL", "No",              "—"],
            ["AWARENESS",    "No",              "—"],
        ],
        col_widths=[1.5, 1.5, 1.5],
    )

    _param_header("Skill Depth", "10 pts", "Resume Auto + BERT")
    _body("BERT classifier assigns depth label per skill. Top 5 skills averaged -> scaled to 10 pts. "
          "Confidence band controls BERT vs evidence blend.")
    _signal_table(
        ["BERT Label",       "Raw Score", "Typical Resume Evidence"],
        [
            ["ARCHITECT_LEVEL", "5.0", "Designed/owned system with this tech; architectural decisions"],
            ["ADVANCED",        "4.0", "Deep project ownership, perf tuning, mentored others on it"],
            ["HANDS_ON",        "3.0", "Actively used in multiple projects"],
            ["FOUNDATIONAL",    "1.5", "Listed but minimal project evidence"],
            ["AWARENESS",       "0.5", "Mentioned in passing, no project usage"],
        ],
        col_widths=[1.4, 0.9, 3.1],
    )
    _signal_table(
        ["Confidence Band",         "Blend Formula"],
        [
            [">=0.65  (high)",       "65% BERT + 35% evidence"],
            ["0.45-0.64  (medium)",  "45% BERT + 55% evidence"],
            ["< 0.45  (low)",        "20% BERT + 80% evidence"],
        ],
        col_widths=[1.8, 2.5],
    )

    _param_header("Skill Recency", "5 pts", "Resume Auto")
    _body("Per skill: recency tag assigned from last job that used it. "
          "Score = sum(weight x recency_factor) / sum(weight) x 5.")
    _signal_table(
        ["Recency Tag", "Definition",                         "Recency Factor"],
        [
            ["CURRENT",  "Used in current / most-recent role",  "1.0"],
            ["RECENT",   "Last used within 2 years",            "1.0"],
            ["MID",      "2-4 years since last used",           "0.5"],
            ["OLD",      "> 4 years since last used",           "0.0"],
        ],
        col_widths=[1.1, 2.5, 1.2],
    )
    _signal_table(
        ["Skill Importance", "Weight Multiplier"],
        [
            ["EXPERT",  "3x"],
            ["DEEP",    "2x"],
            ["APPLIED", "1x"],
        ],
        col_widths=[2.0, 1.5],
    )

    _param_header("Year-on-Year Learning  (skills_learning_acumen)", "5 pts", "Resume Auto")
    _body("Signal: count distinct calendar years with a new skill first introduced.")
    _signal_table(
        ["Condition",                                        "Score"],
        [
            ["fast_learner flag  (new skill added this year)",  "5"],
            ["3+ distinct years with new skills",               "3"],
            ["1-2 years",                                       "2"],
            ["0 years",                                         "0"],
        ],
        col_widths=[3.5, 0.8],
    )

    _param_header("Certifications", "5 pts", "Resume Auto")
    _body("Signal: regex detects cert keywords  (AWS, GCP, Azure, CKA, CKAD, PMP, Databricks, etc.).")
    _signal_table(
        ["Cert Count", "Score"],
        [["4+","5"],["3","4"],["2","2.5"],["1","1"],["0","0"]],
        col_widths=[1.5, 1.0],
    )

    _h3("Recruiter-Stage Skill Parameters  (+7 pts)")

    _param_header("Coding Community  (coding_community)", "4 pts", "Recruiter Stage")
    _body("Resume auto-extracts coding profile URLs  (GitHub, GitLab, Kaggle, LeetCode, Codeforces). "
          "Recruiter validates and assigns 0-4.")
    _signal_table(
        ["Level",    "Evidence",                                                   "Score"],
        [
            ["Elite",   "GitHub 500+ contributions/yr, 5+ repos >=100 stars; Kaggle Master+", "4"],
            ["Strong",  "GitHub 200+ contributions/yr, 2+ repos >=50 stars; Kaggle Expert",   "3"],
            ["Moderate","Active profile, some repos, LeetCode 200+ solved",                   "2"],
            ["Minimal", "Profile exists but low activity",                                    "1"],
            ["None",    "No coding community presence",                                       "0"],
        ],
        col_widths=[1.0, 3.2, 0.6],
    )

    _param_header("Project Explanation  (project_explanation)", "3 pts", "Recruiter Stage")
    _body("STAR-format walk-through quality during phone screen. Recruiter assigns 0-3.")
    _signal_table(
        ["Quality",    "Description",                                    "Score"],
        [
            ["Excellent", "Crisp STAR: summary -> problem -> approach -> metric", "3"],
            ["Good",      "Clear structure, some quantification",               "2"],
            ["Adequate",  "Basic description, no metrics",                      "1"],
            ["Poor",      "Vague or rambling",                                  "0"],
        ],
        col_widths=[1.0, 3.2, 0.6],
    )

    _h3("Panel-Stage Parameters  (+13 pts)")
    _body("All three scored live in the technical interview — resume text cannot improve these.")

    _param_header("Communication Skills", "5 pts", "Panel Stage")
    _body("Verbal + written clarity: technical depth appropriateness, structured thinking, "
          "ability to tailor explanation to audience.")

    _param_header("Domain Skills", "5 pts", "Panel Stage")
    _body("System design depth, algorithms, domain-specific frameworks, "
          "quality of architectural decision-making.")

    _param_header("Problem Solving", "3 pts", "Panel Stage")
    _body("Live coding or case analysis: structured approach, edge-case handling, "
          "time/space complexity awareness.")

    # ══════════════════════════════════════════
    # EDUCATION
    # ══════════════════════════════════════════
    _h2("Education Section — 15 pts")

    _param_header("Institute Tier", "5 pts", "Resume Auto")
    _body("Signal: best institution matched against INSTITUTE_DICTIONARY -> tier classification.")
    _signal_table(
        ["Tier",               "Examples",                                                      "Base", "GPA Bonus"],
        [
            ["TIER_1",          "IIT, IIM, IISc, NIT-T, BITS Pilani; MIT, Stanford, Oxford",     "4.0",  "+1.0 if EXCELLENT"],
            ["TIER_2",          "NITs (non-top), IIIT-H, DCE, VIT, Manipal, good state univs",  "3.0",  "+0.5 if EXCELLENT"],
            ["TIER_3",          "Average private/state colleges",                                "2.0",  "—"],
            ["TIER_4 / Unknown","Unrecognised institution",                                      "1.0",  "—"],
        ],
        col_widths=[1.3, 2.6, 0.7, 1.0],
    )
    _signal_table(
        ["GPA Label",     "Threshold"],
        [
            ["EXCELLENT",    ">=8.5 / 10   or   >=3.7 / 4.0"],
            ["GOOD",         "7.5-8.4 / 10"],
            ["AVERAGE",      "6.0-7.4 / 10"],
            ["BELOW_AVERAGE","< 6.0 / 10"],
        ],
        col_widths=[1.5, 3.0],
    )

    _param_header("Degree Level", "2 pts", "Resume Auto")
    _body("Signal: degree name matched via COURSE_DICTIONARY -> type + IT/non-IT stream.")
    _signal_table(
        ["Degree",                   "Stream",    "Score"],
        [
            ["Ph.D.",                  "IT / CS",   "2.0"],
            ["Ph.D.",                  "Non-IT",    "1.75"],
            ["M.Tech / M.E.",          "IT",        "1.75"],
            ["B.Tech + M.Tech (integ)","IT",        "1.75"],
            ["B.Tech / B.E.",          "IT",        "1.5"],
            ["B.Tech / B.E.",          "Non-IT",    "1.0"],
            ["Diploma",                "Any",       "0.75"],
            ["Not found",              "—",         "1.0  (default)"],
        ],
        col_widths=[2.0, 1.2, 1.1],
    )

    _param_header("Education-Job Relevance", "2 pts", "Resume Auto")
    _body("Signal: degree stream mapped to relevance tier via _STREAM_RELEVANCE lookup.")
    _signal_table(
        ["Relevance",    "Field Examples",                                              "Score"],
        [
            ["HIGH",          "CS, CE, IT, Data Science, AI/ML, ECE, EEE, Software Eng",   "2.0"],
            ["MEDIUM",        "Maths, Statistics, Quant Econ, Analytics, Management",       "1.5"],
            ["FOUNDATIONAL",  "Arts, Humanities, Commerce, non-technical fields",           "0.5"],
            ["UNKNOWN",       "Field not identified from resume text",                      "1.0"],
        ],
        col_widths=[1.3, 2.9, 0.8],
    )

    _param_header("Education Gap", "1 pt", "Resume Auto")
    _body("Signal: months between graduation date and first employment start date.")
    _signal_table(
        ["Gap",           "Score"],
        [
            ["<=6 months",  "1.0"],
            ["6-12 months", "0.5"],
            [">12 months",  "0.0"],
        ],
        col_widths=[2.0, 1.0],
    )

    _param_header("Patents & Publications", "2.5 pts", "Resume Auto")
    _body("Signal: regex detects patent keywords  (patent, filed, granted)  or publication venues "
          "(IEEE, ACM, Springer, NeurIPS, CVPR, arXiv). "
          "TIER_1 gets 0.5 pts baseline; TIER_2 gets 0.25 pts even without explicit mention.")
    _signal_table(
        ["Signal Type",            "Detection Pattern",                    "Score"],
        [
            ["Patent (granted)",     "Granted + patent number present",      "2.5"],
            ["Patent (filed)",       "Filed + application number present",   "2.0"],
            ["Peer-reviewed pub",    "IEEE/ACM/Springer/NeurIPS/CVPR venue", "2.0"],
            ["arXiv / preprint",     "arXiv ID or 'preprint' keyword",       "1.0"],
            ["TIER_1 baseline",      "No explicit mention",                  "0.5"],
            ["TIER_2 baseline",      "No explicit mention",                  "0.25"],
        ],
        col_widths=[1.8, 2.4, 0.8],
    )

    _param_header("Executive Education", "1 pt", "Resume Auto")
    _body("Signal: keywords — 'executive PG', 'online specialisation', 'Coursera', 'upGrad', "
          "'professional diploma', 'distance learning'. Any continuing education beyond primary degree.")

    _param_header("LinkedIn Activity", "1 pt", "Recruiter Stage")
    _body("Recruiter validates: recent posts, recommendation count, endorsements, profile completeness. "
          "3+ recommendations and posts within last 6 months = 1 pt.")

    _param_header("Extra-Curriculars", "1 pt", "Resume Auto")
    _body("Signal: keywords — hackathon, open-source, sports captain, volunteering, NSS, NCC, "
          "cultural events. Any meaningful activity beyond work/academics.")
    _signal_table(
        ["Evidence Type",   "Examples"],
        [
            ["Hackathon",   "Smart India Hackathon, AngelHack, HackMIT — winner/finalist"],
            ["Open Source", "GitHub maintainer, 50+ star repo, OSS contributor"],
            ["Sports",      "Team captain, college sports representative"],
            ["Volunteering","NSS, NGO, teaching, community development"],
        ],
        col_widths=[1.5, 3.5],
    )

    # ══════════════════════════════════════════
    # RECRUITER CHECKLIST
    # ══════════════════════════════════════════
    _h2("Recruiter Quick-Reference Checklist")
    _body("Use before the phone screen — check items visible on the resume:")

    for _sec, _items in [
        ("EXPERIENCE", [
            "10+ years total, continuous (no unexplained gaps)",
            "Clear title progression (at least 2 promotions visible)",
            "Average tenure >= 2 years per role",
            "FAANG / Unicorn / strong product company in history",
            "3+ explicit awards or recognitions",
            "Project 1: quantified impact (% / Rs / $ figure)",
            "Project 2: role + duration + skills stated",
        ]),
        ("SKILLS", [
            "8+ distinct skills listed with years of usage",
            "Primary skill: ARCHITECT or ADVANCED evidence in resume text",
            "Key skills visible in current or last role",
            "New skill added each year for 3+ years",
            "3+ professional certifications with issuer + year",
            "GitHub / Kaggle / LeetCode profile URL present",
        ]),
        ("EDUCATION", [
            "Degree from TIER_1 or TIER_2 institution",
            "CPI / GPA stated  (8+ for TIER_1, 7+ for TIER_2)",
            "CS / CE / IT / DS degree  (HIGH relevance)",
            "First job started within 6 months of graduation",
            "At least one patent or publication",
            "Executive / continuing education program listed",
            "Hackathon or volunteer activity mentioned",
        ]),
        ("RED FLAGS", [
            "3+ companies with avg tenure < 18 months",
            "Gap > 12 months with no explanation",
            "No quantified impact in any project",
            "Skills listed without years of usage",
            "Same job title for 8+ years  (no progression)",
            "No certifications for cloud/infra roles",
        ]),
    ]:
        _h3(_sec)
        for _item in _items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Inches(0.2)
            r = p.add_run(f"{'x' if _sec == 'RED FLAGS' else '[ ]'}  {_item}")
            r.font.size = Pt(9.5)
            if _sec == "RED FLAGS":
                r.font.color.rgb = RED

    # ══════════════════════════════════════════
    # SCORE INTERPRETATION
    # ══════════════════════════════════════════
    _h2("Score Interpretation")
    _signal_table(
        ["Score Range", "Band",    "Action"],
        [
            ["90-100", "STRONG",   "Exceptional. Fast-track to panel."],
            ["75-89",  "GOOD",     "Strong candidate. Proceed with confidence."],
            ["60-74",  "AVERAGE",  "Borderline. Validate gaps in phone screen."],
            ["< 60",   "WEAK",     "Significant gaps. Shortlist only if role is hard to fill."],
        ],
        col_widths=[1.2, 1.0, 3.4],
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF — Perfect Resume  (ReportLab)
# ---------------------------------------------------------------------------

def build_perfect_resume_pdf() -> bytes:
    """Return bytes of a clean PDF version of the 100/100 sample resume."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    PRIMARY = colors.HexColor("#353395")
    MUTED   = colors.HexColor("#62748E")
    BLACK   = colors.HexColor("#262626")
    GREEN   = colors.HexColor("#16A34A")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=1.6*cm, bottomMargin=1.6*cm,
    )

    S = lambda **kw: ParagraphStyle("x", **kw)
    name_s    = S(fontName="Helvetica-Bold", fontSize=22, textColor=PRIMARY, alignment=TA_CENTER, spaceAfter=3)
    contact_s = S(fontName="Helvetica",      fontSize=9,  textColor=MUTED,   alignment=TA_CENTER, spaceAfter=2)
    section_s = S(fontName="Helvetica-Bold", fontSize=8.5,textColor=PRIMARY, spaceBefore=8, spaceAfter=2)
    job_s     = S(fontName="Helvetica-Bold", fontSize=10.5,textColor=BLACK,  spaceBefore=6, spaceAfter=1)
    bullet_s  = S(fontName="Helvetica",      fontSize=9.5, textColor=BLACK,  leftIndent=12, spaceAfter=1)
    skill_s   = S(fontName="Helvetica",      fontSize=9.5, textColor=BLACK,  spaceAfter=2)
    muted_s   = S(fontName="Helvetica-Oblique",fontSize=9, textColor=MUTED,  leftIndent=12, spaceAfter=2)
    edu_s     = S(fontName="Helvetica-Bold", fontSize=10,  textColor=BLACK,  spaceBefore=5, spaceAfter=1)
    sub_s     = S(fontName="Helvetica",      fontSize=9,   textColor=MUTED,  spaceAfter=1)

    def _hr():
        return HRFlowable(width="100%", thickness=1, color=PRIMARY, spaceAfter=3, spaceBefore=1)

    def _section(text):
        return [Paragraph(text.upper(), section_s), _hr()]

    def _job(title, company, period, loc):
        return Paragraph(
            f"<b>{title}</b>  ·  <font color='#353395'><b>{company}</b></font>"
            f"<font color='#62748E' size='9'>  |  {period}  |  {loc}</font>",
            job_s,
        )

    def _b(text):
        return Paragraph(f"• {text}", bullet_s)

    def _sk(label, detail):
        return Paragraph(f"<b><font color='#353395'>{label}:</font></b>  {detail}", skill_s)

    def _edu(degree, inst, year, detail=""):
        items = [Paragraph(f"<b>{degree}</b>  —  <font color='#353395'><b>{inst}</b></font>  ({year})", edu_s)]
        if detail:
            items.append(Paragraph(detail, sub_s))
        return items

    story = []

    # Header
    story += [
        Paragraph("Arjun Mehta", name_s),
        Paragraph("Bangalore, India  ·  arjun.mehta@gmail.com  ·  +91 98765 43210", contact_s),
        Paragraph("linkedin.com/in/arjunmehta  ·  github.com/arjunmehta  ·  arjunmehta.dev", contact_s),
        Spacer(1, 6),
    ]

    # Summary
    story += _section("Professional Summary")
    story += [
        Paragraph(
            "Staff Engineer with 11+ years building high-scale distributed systems across e-commerce, "
            "fintech, and adtech. Led cross-functional teams of 8–12 engineers at FAANG-tier companies. "
            "Architected systems handling 2M+ QPS with sub-10ms p99 latency. "
            "2x patent holder, IEEE published author, CKA + AWS SAP certified.",
            bullet_s,
        ),
        Spacer(1, 4),
    ]

    # Experience
    story += _section("Experience")
    story += [
        _job("Staff Engineer", "Google India", "Jan 2021 - Present (3.5 yr)", "Bangalore"),
        _b("Architected real-time ad-serving pipeline (Go, K8s, Bigtable, Pub/Sub) at 2M QPS, p99 &lt; 8ms — saved $1.4M/yr."),
        _b("Led 8-engineer team; mentored 3 SDEs promoted within 18 months."),
        _b("6-month onsite with Google NYC; presented reliability roadmap to VP Engineering."),
        _b("Won Google 'Peer Bonus' (2022) and 'Spot Award' (2023)."),
        Paragraph("<i><font color='#62748E'>Stack: Go · Kubernetes · Bigtable · Terraform · Prometheus · Pub/Sub · Spanner</font></i>", muted_s),
        _job("Senior Engineer to Staff Engineer", "Swiggy", "Mar 2018 - Dec 2020 (2.8 yr)", "Bangalore"),
        _b("Owned dynamic surge-pricing ML system (Python, XGBoost, Airflow, Kafka) — sole architect; saved Rs 3.2 Cr/month."),
        _b("Built order-matching microservices (FastAPI, Redis, PostgreSQL) for 400K orders/day; reduced latency 28%."),
        _b("Promoted Senior to Staff in 14 months — fastest in cohort. Won 'Engineer of the Year 2020'."),
        Paragraph("<i><font color='#62748E'>Stack: Python · FastAPI · XGBoost · Airflow · Kafka · Redis · PostgreSQL · Docker</font></i>", muted_s),
        _job("Senior Software Engineer", "Flipkart", "Jun 2015 - Feb 2018 (2.8 yr)", "Bangalore"),
        _b("Built product search ranking pipeline (Elasticsearch, Python, Spark) — 18% CTR uplift on 200M SKUs."),
        _b("Conducted 120+ technical interviews; shortlisted 14 hires."),
        Paragraph("<i><font color='#62748E'>Stack: Python · Elasticsearch · Apache Spark · Hive · MySQL · Cassandra</font></i>", muted_s),
        _job("Software Engineer", "Infosys", "Jul 2013 - May 2015 (1.9 yr)", "Pune"),
        _b("Developed REST APIs for a banking client (Java, Spring Boot, Oracle); 0 P1 incidents in 18 months."),
        Paragraph("<i><font color='#62748E'>Stack: Java · Spring Boot · Oracle · Maven · Jenkins</font></i>", muted_s),
    ]

    # Skills
    story += _section("Technical Skills")
    story += [
        _sk("Languages",     "Python (9 yr, Expert)  ·  Go (4 yr, Advanced)  ·  Java (3 yr, Applied)  ·  SQL (10 yr, Expert)"),
        _sk("Frameworks",    "FastAPI (4 yr, Advanced)  ·  Spring Boot (3 yr, Hands-On)  ·  Spark (4 yr, Advanced)"),
        _sk("ML / Data",     "XGBoost (5 yr, Adv)  ·  Scikit-learn (4 yr)  ·  Airflow (4 yr, Adv)  ·  Kafka (5 yr, Adv)"),
        _sk("Infra / Cloud", "Kubernetes (5 yr, Adv)  ·  Terraform (3 yr, Adv)  ·  AWS (4 yr, Adv)  ·  GCP (3 yr)"),
        _sk("Databases",     "PostgreSQL (8 yr, Expert)  ·  Redis (6 yr, Adv)  ·  Bigtable (3 yr)  ·  Cassandra (3 yr)"),
        _sk("Coding",        "github.com/arjunmehta  ·  LeetCode 350+ solved  ·  Kaggle Expert"),
    ]

    # Certifications
    story += _section("Certifications")
    for cert, issuer, yr in [
        ("AWS Solutions Architect Professional", "Amazon Web Services", "2023"),
        ("Certified Kubernetes Administrator (CKA)", "CNCF", "2022"),
        ("Google Professional Machine Learning Engineer", "Google Cloud", "2021"),
        ("Databricks Certified Associate Developer for Apache Spark", "Databricks", "2020"),
    ]:
        story.append(Paragraph(f"• <b>{cert}</b>  —  <font color='#62748E'>{issuer}  ({yr})</font>", bullet_s))

    # Education
    story += _section("Education")
    story += _edu("M.Tech Computer Science", "IIT Bombay", "2011-2013",
                   "CPI: 9.1 / 10  ·  Thesis: 'Distributed consensus optimisation in heterogeneous networks'")
    story += _edu("B.Tech Computer Engineering", "IIT Bombay", "2007-2011",
                   "CPI: 8.7 / 10  ·  Institute Silver Medal for academic excellence")

    # Publications & Patents
    story += _section("Publications & Patents")
    for pub in [
        "Mehta A. et al. — 'Adaptive Latency-Aware Scheduling in Real-Time Bidding Systems', IEEE INFOCOM 2023",
        "Mehta A., Kumar R. — 'Low-Latency Feature Store for Online ML Inference', KDD Workshop 2021",
        "Patent: 'Method and System for Dynamic Surge Pricing in Last-Mile Logistics' (IN Patent 412763, Filed 2020, Granted 2022)",
    ]:
        story.append(Paragraph(f"• {pub}", bullet_s))

    # Achievements
    story += _section("Achievements & Activities")
    for ex in [
        "Smart India Hackathon 2012 — National Winner (Team of 4)",
        "Open-source contributor: 3 repos with 800+ GitHub stars; maintainer of 'pystream-ml' library",
        "Executive PG in AI &amp; ML — upGrad / IIIT Bangalore (2020, Distinction)",
        "NSS Volunteer (2008-2011) — taught coding to 200+ school students",
        "Chess club captain, IIT Bombay (2009-2010)",
    ]:
        story.append(Paragraph(f"• {ex}", bullet_s))

    doc.build(story)
    return buf.getvalue()