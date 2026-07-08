"""
Tvarah Portal REST API  (v1)
============================
JSON-only FastAPI router mounted at /api/v1 in app.py.
Provides JWT-auth endpoints for the tvarah-portal-ui frontend.

Auth flow:
  POST /api/v1/auth/login  → {access_token, token_type}
  GET  /api/v1/auth/me     → user profile

Candidate APIs:
  GET    /api/v1/candidates          → paginated list with AI scores
  GET    /api/v1/candidates/{id}     → full profile
  POST   /api/v1/candidates          → create candidate record
  PATCH  /api/v1/candidates/{id}     → update stage / status / assignment
  POST   /api/v1/candidates/{id}/analyze   → trigger AI analysis
  GET    /api/v1/candidates/{id}/analysis  → get latest AI analysis
  POST   /api/v1/candidates/{id}/outcome   → record placement outcome
  GET    /api/v1/candidates/{id}/outcome   → get outcome

Job (JD) APIs:
  GET    /api/v1/jobs                → list JDs
  POST   /api/v1/jobs                → create JD
  GET    /api/v1/jobs/{id}           → JD detail
  PATCH  /api/v1/jobs/{id}           → update JD
  DELETE /api/v1/jobs/{id}           → close JD
  POST   /api/v1/jobs/{id}/match_all → run AI match for all candidates

Matching:
  GET    /api/v1/candidates/{cid}/match/{jd_id}  → match result
  POST   /api/v1/candidates/{cid}/match/{jd_id}  → trigger match

Dashboard:
  GET    /api/v1/dashboard/stats     → aggregate stats

Users:
  GET    /api/v1/users               → list portal users
  POST   /api/v1/users/invite        → create user + set temp password
  PATCH  /api/v1/users/{id}          → update role / status

Skills reference:
  GET    /api/v1/ref/skills          → all skills
  GET    /api/v1/ref/job_titles      → all job titles
"""
from __future__ import annotations

import json
import logging
import os
import secrets
import smtplib
import uuid
from contextlib import contextmanager
from datetime import datetime, timedelta, timezone
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from typing import Any, Generator

import bcrypt
import jwt
import psycopg2
import psycopg2.extras
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile, status
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer

logger = logging.getLogger("portal_api")

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
_JWT_SECRET = os.getenv("PORTAL_JWT_SECRET", "tvarah-portal-secret-change-in-prod-2025")
_JWT_ALGO = "HS256"
_JWT_EXPIRE_HOURS = 24 * 7  # 7 days

# ---------------------------------------------------------------------------
# SMTP / OTP config
# ---------------------------------------------------------------------------
_SMTP_HOST = "smtp.office365.com"
_SMTP_PORT = 587
_SMTP_USER = os.getenv("SMTP_EMAIL", "admin@tvarah.com")
_SMTP_PASS = os.getenv("SMTP_PASSWORD", "")
_OTP_TTL_MINUTES = 10
# Dev mode: true when SMTP password is absent or placeholder — OTP returned in API response
_DEV_MODE = not _SMTP_PASS or _SMTP_PASS in ("PLACEHOLDER_SET_ME", "placeholder")

# In-memory OTP store: email → (code, expires_at)
_otp_store: dict[str, tuple[str, datetime]] = {}


def _send_otp_email(to_email: str, otp: str) -> None:
    msg = MIMEMultipart("alternative")
    msg["From"]    = _SMTP_USER
    msg["To"]      = to_email
    msg["Subject"] = "Your Tvarah login code"

    text = (
        f"Your Tvarah Portal login code is: {otp}\n\n"
        f"This code is valid for {_OTP_TTL_MINUTES} minutes.\n"
        "If you did not request this, ignore this email."
    )
    html = f"""
<div style="font-family:sans-serif;max-width:420px;margin:0 auto">
  <h2 style="color:#111">Your Tvarah login code</h2>
  <p style="font-size:36px;font-weight:bold;letter-spacing:8px;color:#4f46e5">{otp}</p>
  <p style="color:#555">Valid for {_OTP_TTL_MINUTES} minutes. Do not share this code.</p>
  <hr style="border:none;border-top:1px solid #eee;margin:24px 0">
  <p style="color:#999;font-size:12px">Tvarah Recruitment Platform</p>
</div>"""

    msg.attach(MIMEText(text, "plain"))
    msg.attach(MIMEText(html, "html"))

    with smtplib.SMTP(_SMTP_HOST, _SMTP_PORT, timeout=15) as server:
        server.ehlo()
        server.starttls()
        server.login(_SMTP_USER, _SMTP_PASS)
        server.send_message(msg)


_PG = {
    "host":     os.environ.get("DB_HOST", "10.0.0.6"),
    "port":     int(os.environ.get("DB_PORT", "5432")),
    "dbname":   os.environ.get("DB_NAME", "tvarah"),
    "user":     os.environ.get("DB_USER", "tvarah"),
    "password": os.environ.get("DB_PASSWORD", "4pVhZ873Rm6C"),
}

# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------

@contextmanager
def _db() -> Generator:
    conn = psycopg2.connect(**_PG, cursor_factory=psycopg2.extras.RealDictCursor)
    cur = conn.cursor()
    try:
        yield cur
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def _row(cur) -> dict | None:
    r = cur.fetchone()
    return dict(r) if r else None


def _rows(cur) -> list[dict]:
    return [dict(r) for r in (cur.fetchall() or [])]


# ---------------------------------------------------------------------------
# JWT helpers
# ---------------------------------------------------------------------------

def _make_token(user_id: str, email: str, role: str) -> str:
    exp = datetime.now(timezone.utc) + timedelta(hours=_JWT_EXPIRE_HOURS)
    payload = {"sub": user_id, "email": email, "role": role, "exp": exp}
    return jwt.encode(payload, _JWT_SECRET, algorithm=_JWT_ALGO)


def _make_keycloak_token(user_id: str, email: str, name: str, role: str) -> str:
    """JWT with Keycloak-compatible claims consumed by tvarah-portal-ui auth.ts."""
    now = datetime.now(timezone.utc)
    exp = now + timedelta(hours=_JWT_EXPIRE_HOURS)
    payload = {
        "sub": user_id,
        "email": email,
        "name": name,
        "iat": int(now.timestamp()),
        "exp": exp,
        "realm_access": {"roles": [role]},
        "resource_access": {
            "tvarahbackend": {"roles": [role]},
            "tvarah-portal":  {"roles": [role]},
        },
    }
    return jwt.encode(payload, _JWT_SECRET, algorithm=_JWT_ALGO)


def _decode_token(token: str) -> dict:
    try:
        return jwt.decode(token, _JWT_SECRET, algorithms=[_JWT_ALGO])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


_bearer = HTTPBearer(auto_error=False)


def get_current_user(
    creds: HTTPAuthorizationCredentials | None = Depends(_bearer),
) -> dict:
    if not creds:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return _decode_token(creds.credentials)


def require_admin(user: dict = Depends(get_current_user)) -> dict:
    if user.get("role") not in ("super_admin", "Site Admin", "admin"):
        raise HTTPException(status_code=403, detail="Admin only")
    return user


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------
router = APIRouter(prefix="/api/v1", tags=["Portal API"])


# ── Auth ─────────────────────────────────────────────────────────────────────

@router.post("/auth/login")
def login(payload: dict) -> dict:
    """Authenticate and return a JWT.

    Body: {email, password}
    """
    email = (payload.get("email") or "").strip().lower()
    password = payload.get("password") or ""
    if not email or not password:
        raise HTTPException(400, "email and password required")

    with _db() as cur:
        cur.execute(
            'SELECT id, name, email, role, password FROM users WHERE LOWER(email)=%s',
            (email,),
        )
        user = _row(cur)

    if not user:
        raise HTTPException(401, "Invalid credentials")
    if not user.get("password"):
        raise HTTPException(401, "Password not set — contact admin")
    if not bcrypt.checkpw(password.encode(), user["password"].encode()):
        raise HTTPException(401, "Invalid credentials")

    name_parts = (user.get("name") or email.split("@")[0]).split(" ", 1)
    token = _make_token(str(user["id"]), user["email"], user["role"] or "recruiter")
    return {
        "access_token": token,
        "token_type": "Bearer",
        "user": {
            "id": str(user["id"]),
            "first_name": name_parts[0],
            "last_name": name_parts[1] if len(name_parts) > 1 else "",
            "email": user["email"],
            "role": user["role"] or "recruiter",
        },
    }


@router.post("/auth/send-otp")
def send_otp(payload: dict) -> dict:
    """Step 1 of email-OTP login: generate a 6-digit code and email it.

    Body: {email}
    """
    email = (payload.get("email") or "").strip().lower()
    if not email:
        raise HTTPException(400, "email required")

    # Validate the user exists
    with _db() as cur:
        cur.execute('SELECT id FROM users WHERE LOWER(email)=%s', (email,))
        user = _row(cur)
    if not user:
        # Return 200 even on unknown email to avoid user enumeration
        return {"status": "success", "message": "If that email is registered, an OTP will be sent"}

    # Generate and store OTP
    code = str(secrets.randbelow(900000) + 100000)  # 100000–999999
    _otp_store[email] = (code, datetime.now(timezone.utc) + timedelta(minutes=_OTP_TTL_MINUTES))

    if _DEV_MODE:
        logger.warning("[DEV] OTP for %s = %s", email, code)
        return {"status": "success", "message": "Dev mode — OTP auto-filled below", "dev_otp": code}

    try:
        _send_otp_email(email, code)
    except Exception as exc:
        logger.error("Failed to send OTP email to %s: %s", email, exc)
        raise HTTPException(500, "Failed to send OTP email — contact admin")

    return {"status": "success", "message": f"OTP sent to {email}"}


@router.post("/auth/verify-otp")
def verify_otp(payload: dict) -> dict:
    """Step 2 of email-OTP login: verify the 6-digit code and return a JWT.

    Returns a Keycloak-compatible token envelope so next-auth can parse roles.
    Body: {email, otp}
    """
    email = (payload.get("email") or "").strip().lower()
    otp   = (payload.get("otp") or "").strip()
    if not email or not otp:
        raise HTTPException(400, "email and otp required")

    # Validate OTP from store
    entry = _otp_store.get(email)
    if not entry:
        raise HTTPException(401, "No OTP requested for this email — send OTP first")
    stored_code, expires_at = entry
    if datetime.now(timezone.utc) > expires_at:
        _otp_store.pop(email, None)
        raise HTTPException(401, "OTP expired — request a new one")
    if stored_code != otp:
        raise HTTPException(401, "Invalid OTP")
    _otp_store.pop(email, None)  # single-use

    # Get user from DB
    with _db() as cur:
        cur.execute(
            'SELECT id, name, email, role FROM users WHERE LOWER(email)=%s',
            (email,),
        )
        user = _row(cur)

    if not user:
        raise HTTPException(401, "User not found")

    name = user.get("name") or email.split("@")[0]
    access_token = _make_keycloak_token(str(user["id"]), user["email"], name, user["role"] or "recruiter")

    return {
        "status": "success",
        "data": {
            "token": {
                "access_token":  access_token,
                "refresh_token": "",
                "expires_in":    _JWT_EXPIRE_HOURS * 3600,
            },
            "firstTimeUser": False,
        },
    }


@router.get("/auth/me")
def me(user: dict = Depends(get_current_user)) -> dict:
    with _db() as cur:
        cur.execute(
            'SELECT id, name, email, role FROM users WHERE id=%s',
            (user["sub"],),
        )
        profile = _row(cur)
    if not profile:
        raise HTTPException(404, "User not found")
    return profile


# ── JD Parse (called by portal's MlJDParser via JD_PARSER_URL) ───────────────

@router.post("/jdParse")
async def jd_parse(file: UploadFile = File(None), text: str = Form(None)) -> dict:
    """Parse a job description PDF/DOCX or raw text using the LLM.

    Returns MlJDResponse format expected by tvarah-portal-ui's jd-parser.ts.
    """
    from llm_client import call_llm_json, analysis_model

    raw_text = ""

    if file and file.size > 0:
        content = await file.read()
        fname = (file.filename or "file.pdf").lower()
        if fname.endswith(".docx"):
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(content))
                raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                raw_text = content.decode("utf-8", errors="ignore")
        else:
            try:
                import pdfplumber, io
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    raw_text = "\n".join(
                        page.extract_text() or "" for page in pdf.pages
                    )
            except Exception:
                try:
                    import pypdf, io as _io
                    reader = pypdf.PdfReader(_io.BytesIO(content))
                    raw_text = "\n".join(
                        (page.extract_text() or "") for page in reader.pages
                    )
                except Exception:
                    raw_text = content.decode("utf-8", errors="ignore")

    elif text:
        raw_text = text

    if not raw_text.strip():
        raise HTTPException(400, "Could not extract text from the file")

    prompt = f"""You are a precise job description parser. Extract all information from the JD below and return ONLY a valid JSON object with this exact structure (no markdown, no extra keys):

{{
  "jd_data": {{
    "role_title": "string",
    "company_name": "string or null",
    "salary_range": "string or null",
    "location": [{{"city": "string or null", "country": "string or null"}}],
    "job_type": "Full Time|Part Time|Contract|Internship",
    "job_level": "Junior|Mid|Senior|Lead|Manager",
    "work_mode": "Onsite|Remote|Hybrid",
    "mandatory_skills": {{
      "programming_languages": [],
      "frameworks_and_libraries": [],
      "tools": [],
      "databases": [],
      "cloud_and_infra": []
    }},
    "optional_skills": [],
    "min_years_experience": null,
    "max_years_experience": null,
    "degree_required": "string or null",
    "summary_responsibilities": []
  }}
}}

Job Description:
---
{raw_text[:6000]}
---"""

    model = analysis_model("z-ai/glm-4.5-air:free")
    result = call_llm_json(model, [{"role": "user", "content": prompt}], max_tokens=1200)
    if not result or "jd_data" not in result:
        raise HTTPException(500, "LLM failed to parse job description")
    return result


# ── Resume Parse (called by portal's MlResumeParser via RESUME_PARSER_URL) ───

@router.post("/resumeParse")
async def resume_parse(file: UploadFile = File(None)) -> dict:
    """Parse a resume PDF/DOCX into MlResumeResponseV2 format expected by tvarah-portal-ui."""
    from llm_client import call_llm_json, analysis_model
    import io

    if not file or file.size == 0:
        raise HTTPException(400, "No file provided")

    content = await file.read()
    fname = (file.filename or "resume.pdf").lower()
    raw_text = ""

    if fname.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            raw_text = content.decode("utf-8", errors="ignore")
    else:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                raw_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception:
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                raw_text = "\n".join((page.extract_text() or "") for page in reader.pages)
            except Exception:
                raw_text = content.decode("utf-8", errors="ignore")

    if not raw_text.strip():
        raise HTTPException(400, "No text found in file")

    # ── Pre-extract name fallback ──────────────────────────────────────────────
    # Heuristic: candidate name is usually the very first non-empty line of a resume.
    _fallback_name = ""
    for _line in raw_text.splitlines():
        _stripped = _line.strip()
        # Skip short lines, all-caps headers, or lines with @ / digits (email/phone)
        if (len(_stripped) > 3 and not _stripped.isupper()
                and "@" not in _stripped and not _stripped[0].isdigit()):
            _fallback_name = _stripped
            break

    # Also try to extract from filename: "ShashankSrivastava_DS.pdf" → "Shashank Srivastava"
    if not _fallback_name:
        import re as _re
        _stem = (file.filename or "").rsplit(".", 1)[0]  # drop extension
        _stem = _re.split(r"[_\-\s]", _stem)[0]         # take first segment
        # Insert spaces before capital letters: "ShashankSrivastava" → "Shashank Srivastava"
        _fallback_name = _re.sub(r"(?<=[a-z])(?=[A-Z])", " ", _stem)

    # ── Two-step LLM parse: basic info first, then experience/skills ─────────
    prompt_basic = f"""Extract candidate information from this resume. Return ONLY valid JSON, no markdown.

{{
  "full_name": "candidate full name",
  "first_name": "first name",
  "last_name": "last name",
  "email": "email or null",
  "phone": "phone or null",
  "city": "current city or null",
  "summary": "2 sentence candidate summary"
}}

RESUME (first 1500 chars):
{raw_text[:1500]}"""

    prompt_details = f"""Extract work experience, education and skills from this resume. Return ONLY valid JSON, no markdown.

{{
  "work_experience": [
    {{"company": "name", "title": "job title", "start": "YYYY-MM", "end": "YYYY-MM or null", "current": false, "description": "key responsibilities"}}
  ],
  "education": [
    {{"institution": "name", "degree": "B.Tech etc", "field": "subject", "start": "YYYY-MM", "end": "YYYY-MM or null", "current": false}}
  ],
  "skills": {{
    "programming_languages": [],
    "frameworks_and_libraries": [],
    "tools_and_platforms": [],
    "databases": [],
    "cloud_and_infra": [],
    "soft_skills": []
  }}
}}

RESUME:
{raw_text[:4000]}"""

    model = analysis_model("z-ai/glm-4.5-air:free")
    try:
        basic  = call_llm_json(model, [{"role": "user", "content": prompt_basic}],   max_tokens=600)
    except Exception as _e:
        logger.warning("resume_parse basic LLM error: %s", _e)
        basic = None
    try:
        detail = call_llm_json(model, [{"role": "user", "content": prompt_details}], max_tokens=2000)
    except Exception as _e:
        logger.warning("resume_parse detail LLM error: %s", _e)
        detail = None

    if not basic:
        raise HTTPException(500, "LLM failed to parse resume (basic info step failed)")

    detail = detail or {}

    # ── Assemble MlResumeResponseV2 ────────────────────────────────────────
    full_name = (basic.get("full_name") or "").strip()
    result = {
        "basic_info": {
            "personal_info": {
                "full_name":        full_name,
                "first_name":       basic.get("first_name") or "",
                "middle_name":      None,
                "last_name":        basic.get("last_name") or "",
                "date_of_birth":    None,
                "gender":           None,
                "nationality":      None,
                "work_authorization": None,
            },
            "contact_info": {
                "primary_email":        basic.get("email"),
                "secondary_email":      None,
                "primary_phone_number": basic.get("phone"),
                "secondary_phone_number": None,
                "country_code":         None,
                "current_city":         basic.get("city"),
                "current_state":        None,
                "current_country":      None,
                "linkedin":             None,
                "github":               None,
                "portfolio":            None,
            },
        },
        "insight_info": {
            "education_info": [
                {
                    "institution_name": e.get("institution", ""),
                    "degree":           e.get("degree", ""),
                    "field_of_study":   e.get("field", ""),
                    "specialisation":   None,
                    "education_level":  "Graduate",
                    "start_date":       e.get("start"),
                    "end_date":         e.get("end"),
                    "is_current":       "Yes" if e.get("current") else "No",
                    "grade_or_gpa":     None,
                    "mode":             "Full-time",
                }
                for e in (detail.get("education") or [])
            ],
            "education_insights": "",
            "work_experience_info": [
                {
                    "company_name":      w.get("company", ""),
                    "company_location":  None,
                    "job_title":         w.get("title", ""),
                    "employment_type":   "Full-time",
                    "start_date":        w.get("start"),
                    "end_date":          w.get("end"),
                    "is_current_role":   "Yes" if w.get("current") else "No",
                    "role_description":  w.get("description", ""),
                    "experience_insights": None,
                }
                for w in (detail.get("work_experience") or [])
            ],
            "skills_info": detail.get("skills") or {
                "programming_languages": [], "frameworks_and_libraries": [],
                "tools_and_platforms": [], "databases": [],
                "cloud_and_infra": [], "soft_skills": [], "certified_skills": [],
            },
            "overall_insights": basic.get("summary", ""),
        },
        "resume_score": None,
        "intelligence_insights": None,
        "strengths_weaknesses": None,
        "tagging": None,
        "skill_depth_analysis": None,
    }

    # ── Ensure full_name is never empty ───────────────────────────────────────
    pi = result.get("basic_info", {}).get("personal_info", {})
    if not pi.get("full_name", "").strip() and _fallback_name:
        pi["full_name"] = _fallback_name
        # Also split into first/last if missing
        parts = _fallback_name.split()
        if not pi.get("first_name"):
            pi["first_name"] = parts[0]
        if not pi.get("last_name") and len(parts) > 1:
            pi["last_name"] = parts[-1]

    return result


# ── Job Descriptions (sync endpoint called by portal's create-backend route) ──
# Uses the existing job_postings table in PostgreSQL.
# Columns job_type, job_level, work_mode, role_title were added via ALTER TABLE.

@router.get("/job-descriptions")
def list_job_descriptions(user: dict = Depends(get_current_user)) -> dict:
    """Return all portal-uploaded JDs from job_postings in standard envelope format."""
    import json as _json

    with _db() as cur:
        cur.execute(
            """SELECT jd_id, COALESCE(role_title, title) AS role_title,
                      company, mandatory_skills, nice_to_have,
                      COALESCE(job_type, 'Full Time') AS job_type,
                      COALESCE(job_level, 'Mid') AS job_level,
                      COALESCE(work_mode, 'Onsite') AS work_mode,
                      yoe_min, yoe_max, jd_status, created_at
               FROM job_postings
               WHERE jd_status != 'closed'
               ORDER BY created_at DESC NULLS LAST"""
        )
        rows = _rows(cur)

    def _parse_skills(raw) -> list:
        if not raw:
            return []
        if isinstance(raw, list):
            return raw
        try:
            parsed = _json.loads(raw)
            if isinstance(parsed, list):
                return parsed
            return []
        except Exception:
            return []

    def _status_map(s: str) -> str:
        return {"active": "Open", "open": "Open", "closed": "Closed", "draft": "Draft"}.get(
            (s or "active").lower(), "Open"
        )

    jds = []
    for r in rows:
        jds.append({
            "id":               r["jd_id"],
            "code":             r["jd_id"],
            "jobTitle":         r.get("role_title") or "",
            "company":          r.get("company") or "",
            "status":           _status_map(r.get("jd_status") or "active"),
            "jobType":          r.get("job_type") or "Full Time",
            "jobMode":          r.get("work_mode") or "Onsite",
            "jobLevel":         r.get("job_level") or "Mid",
            "experienceMinYrs": r.get("yoe_min"),
            "experienceMaxYrs": r.get("yoe_max"),
            "requiredSkills":   _parse_skills(r.get("mandatory_skills")),
            "goodToHaveSkills": _parse_skills(r.get("nice_to_have")),
            "createdAt":        r["created_at"] if isinstance(r.get("created_at"), str)
                                else (r["created_at"].isoformat() if r.get("created_at") else None),
            "totalPositions":   1,
            "totalPositionsSelected": 0,
        })

    return {"status": "success", "message": "ok", "data": jds}


@router.post("/job-descriptions", status_code=201)
def create_job_description(payload: dict, user: dict = Depends(get_current_user)) -> dict:
    """Receive a parsed JD from the portal and save it to job_postings table."""
    title   = (payload.get("jobTitle") or "").strip()
    if not title:
        raise HTTPException(400, "jobTitle required")
    company = (payload.get("company") or "Unknown").strip()
    jd_id   = f"JD-{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}"
    import json as _json
    now_str = datetime.now(timezone.utc).isoformat()
    with _db() as cur:
        cur.execute(
            """INSERT INTO job_postings
               (jd_id, title, role_title, company, mandatory_skills, nice_to_have,
                job_type, job_level, work_mode, yoe_min, yoe_max, jd_status, created_at)
               VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'active',%s)
               ON CONFLICT (jd_id) DO NOTHING""",
            (
                jd_id, title, title, company,
                _json.dumps(payload.get("requiredSkills", [])),
                _json.dumps(payload.get("goodToHaveSkills", [])),
                payload.get("jobType", "Full Time"),
                payload.get("jobLevel", "Mid"),
                payload.get("jobMode", "Onsite"),
                payload.get("experienceMinYrs"),
                payload.get("experienceMaxYrs"),
                now_str,
            ),
        )
    return {"ok": True, "jd_id": jd_id}


# ── Standup ───────────────────────────────────────────────────────────────────

@router.get("/standup/board")
def standup_board(user: dict = Depends(get_current_user)) -> dict:
    """Full team standup board: each Prisma user with their JDs, candidate stage counts, and today's standup."""
    today = datetime.now(timezone.utc).date().isoformat()
    with _db() as cur:
        # Team from Prisma users table
        cur.execute(
            "SELECT id, name, email, role FROM users WHERE role NOT IN ('company','client_login') ORDER BY name"
        )
        team = _rows(cur)

        # Today's standup entries (user_id = Prisma user cuid stored by portal)
        cur.execute(
            "SELECT user_id, today, blockers, priorities, created_at FROM standup_updates WHERE date=%s",
            (today,),
        )
        standup_map = {r["user_id"]: r for r in _rows(cur)}

        # JD assignments + candidate stage counts from Prisma positions/candidates tables
        try:
            cur.execute("""
                SELECT
                    p."assignedTo"    AS assigned_to,
                    p."jdId"          AS jd_id,
                    p.title,
                    p.status,
                    p.company,
                    COUNT(c.id)                                                   AS total,
                    COUNT(c.id) FILTER (WHERE c.stage = 'Telephonic')             AS telephonic,
                    COUNT(c.id) FILTER (WHERE c.stage = 'Panel')                  AS panel,
                    COUNT(c.id) FILTER (WHERE c.stage = 'Client')                 AS client_stage,
                    COUNT(c.id) FILTER (WHERE c."poolReady" = TRUE)               AS pool_ready
                FROM positions p
                LEFT JOIN candidates c ON c."jdId" = p."jdId"
                WHERE p."assignedTo" IS NOT NULL
                GROUP BY p."assignedTo", p."jdId", p.title, p.status, p.company
                ORDER BY p."createdAt" DESC
            """)
            jd_map: dict = {}
            for row in _rows(cur):
                at = row.get("assigned_to")
                if at:
                    jd_map.setdefault(at, []).append(row)
        except Exception:
            jd_map = {}

    board = []
    for m in team:
        uid = m["id"]
        jds = jd_map.get(uid, [])
        total_candidates = sum(j.get("total", 0) or 0 for j in jds)
        board.append({
            "user_id":          uid,
            "name":             m["name"],
            "email":            m["email"],
            "role":             m["role"],
            "standup_today":    standup_map.get(uid),
            "jds":              jds,
            "jd_count":         len(jds),
            "total_candidates": total_candidates,
        })

    return {"board": board, "date": today}


@router.get("/standup/team")
def standup_team(limit: int = Query(50, le=200), user: dict = Depends(get_current_user)) -> dict:
    """Return recent standup entries for the whole team."""
    from database import list_standups_for_team
    # Also try to join against Prisma users for name resolution
    today_only = False
    rows = list_standups_for_team(limit=limit)
    # Enrich with Prisma user names where app_users join misses
    missing_names = [r for r in rows if not r.get("full_name")]
    if missing_names:
        ids = list({r["user_id"] for r in missing_names})
        with _db() as cur:
            cur.execute(
                f"SELECT id, name, email FROM users WHERE id = ANY(%s::text[])",
                (ids,),
            )
            prisma_users = {r["id"]: r for r in _rows(cur)}
        for r in rows:
            if not r.get("full_name"):
                pu = prisma_users.get(r["user_id"])
                if pu:
                    r["full_name"] = pu["name"]
                    r["email"] = pu.get("email", r.get("email"))
    return {"standups": rows}


@router.get("/standup/today")
def standup_today(user: dict = Depends(get_current_user)) -> dict:
    """Return today's standup entries for the calling user."""
    from database import get_standup_by_date
    today = datetime.now(timezone.utc).date().isoformat()
    user_id = user.get("sub", "")
    rows = get_standup_by_date(user_id, today)
    return {"standups": rows}


@router.post("/standup")
def submit_standup(payload: dict, user: dict = Depends(get_current_user)) -> dict:
    """Submit today's standup for the calling user."""
    from database import upsert_standup
    today = datetime.now(timezone.utc).date().isoformat()
    user_id = user.get("sub", "")
    if not user_id:
        raise HTTPException(400, "Cannot determine user ID from token")
    update_id = upsert_standup(user_id, today, {
        "jd_id":      payload.get("jd_id"),
        "today":      payload.get("today", ""),
        "blockers":   payload.get("blockers", ""),
        "priorities": payload.get("priorities", ""),
    })
    return {"ok": True, "update_id": update_id}


# ── Candidates ────────────────────────────────────────────────────────────────

@router.get("/candidates")
def list_candidates(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    search: str = Query("", alias="q"),
    stage: str = Query(""),
    role_family: str = Query(""),
    min_score: float = Query(None),
    user: dict = Depends(get_current_user),
) -> dict:
    """List candidates with their AI scores and basic info."""
    offset = (page - 1) * page_size
    filters: list[str] = []
    params: list[Any] = []

    if search:
        filters.append(
            "(LOWER(c.first_name || ' ' || c.last_name) LIKE %s "
            "OR LOWER(c.primary_email) LIKE %s)"
        )
        like = f"%{search.lower()}%"
        params += [like, like]
    if stage:
        filters.append("COALESCE(ap.stage, 'SOURCED') = %s")
        params.append(stage)
    if role_family:
        filters.append("ap.role_family = %s")
        params.append(role_family)
    if min_score is not None:
        filters.append("COALESCE(ap.combined_score, 0) >= %s")
        params.append(min_score)

    where = ("WHERE " + " AND ".join(filters)) if filters else ""

    with _db() as cur:
        cur.execute(
            f"""SELECT COUNT(*) as total
                FROM candidate c
                LEFT JOIN candidate_ai_profile ap ON ap.candidate_id = c.id
                {where}""",
            params,
        )
        total = _row(cur)["total"]

        cur.execute(
            f"""SELECT
                    c.id, c.first_name, c.last_name, c.primary_email,
                    c.total_experience, c.relevant_experience, c.status,
                    ap.resume_score, ap.recruiter_score, ap.panel_score,
                    ap.combined_score, ap.band, ap.role_family, ap.dna_fit,
                    ap.stage, ap.pipeline_status, ap.updated_at,
                    jt.name as job_title
                FROM candidate c
                LEFT JOIN candidate_ai_profile ap ON ap.candidate_id = c.id
                LEFT JOIN job_title jt ON jt.id = c.job_title_id
                {where}
                ORDER BY ap.combined_score DESC NULLS LAST, c.created_on DESC
                LIMIT %s OFFSET %s""",
            params + [page_size, offset],
        )
        rows = _rows(cur)

    # Serialize UUIDs
    candidates = []
    for r in rows:
        d = {k: str(v) if isinstance(v, uuid.UUID) else v for k, v in r.items()}
        candidates.append(d)

    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": (total + page_size - 1) // page_size,
        "candidates": candidates,
    }


@router.get("/candidates/{candidate_id}")
def get_candidate(candidate_id: str, user: dict = Depends(get_current_user)) -> dict:
    """Full candidate profile: base info + experience + education + skills + AI profile."""
    with _db() as cur:
        # Base
        cur.execute(
            """SELECT c.*, jt.name as job_title_name, co.name as current_company_name
               FROM candidate c
               LEFT JOIN job_title jt ON jt.id = c.job_title_id
               LEFT JOIN company co ON co.id = c.current_company_id
               WHERE c.id = %s::uuid""",
            (candidate_id,),
        )
        base = _row(cur)
        if not base:
            raise HTTPException(404, "Candidate not found")

        # Experience
        cur.execute(
            """SELECT ce.*, co.name as company_name, jt.name as job_title_name,
                      d.name as designation_name
               FROM candidate_experience ce
               LEFT JOIN company co ON co.id = ce.company_id
               LEFT JOIN job_title jt ON jt.id = ce.job_title_id
               LEFT JOIN designation d ON d.id = ce.designation_id
               WHERE ce.candidate_id = %s::uuid
               ORDER BY ce.start_date DESC NULLS LAST""",
            (candidate_id,),
        )
        experience = _rows(cur)

        # Education
        cur.execute(
            """SELECT edu.*, inst.name as institution_name,
                      deg.name as degree_name, sp.name as specialization_name
               FROM candidate_education edu
               LEFT JOIN institution inst ON inst.id = edu.institution_id
               LEFT JOIN degree deg ON deg.id = edu.degree_id
               LEFT JOIN specialization sp ON sp.id = edu.specialization_id
               WHERE edu.candidate_id = %s::uuid
               ORDER BY edu.start_date DESC NULLS LAST""",
            (candidate_id,),
        )
        education = _rows(cur)

        # Skills
        cur.execute(
            """SELECT cs.*, sk.name as skill_name, sk.category, sk.tier
               FROM candidate_skill cs
               LEFT JOIN skill sk ON sk.id = cs.skill_id
               WHERE cs.candidate_id = %s::uuid
               ORDER BY cs.is_primary DESC, cs.score DESC NULLS LAST""",
            (candidate_id,),
        )
        skills = _rows(cur)

        # AI profile
        cur.execute(
            "SELECT * FROM candidate_ai_profile WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        ai_profile = _row(cur)

        # Resume summary (BERT insights)
        cur.execute(
            "SELECT * FROM candidate_resume_summary WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        resume_summary = _row(cur)

        # Intelligence insight
        cur.execute(
            "SELECT * FROM candidate_intelligence_insight WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        intelligence = _row(cur)

        # Score
        cur.execute(
            "SELECT * FROM candidate_score WHERE candidate_id = %s::uuid ORDER BY created_on DESC LIMIT 1",
            (candidate_id,),
        )
        score = _row(cur)

        # Outcome
        cur.execute(
            "SELECT * FROM candidate_outcome WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        outcome = _row(cur)

    def _ser(d):
        if d is None:
            return None
        result = {}
        for k, v in d.items():
            if isinstance(v, uuid.UUID):
                result[k] = str(v)
            elif isinstance(v, (datetime,)):
                result[k] = v.isoformat()
            elif isinstance(v, memoryview):
                result[k] = None  # skip binary blobs
            else:
                result[k] = v
        return result

    return {
        "candidate": _ser(base),
        "experience": [_ser(e) for e in experience],
        "education": [_ser(e) for e in education],
        "skills": [_ser(s) for s in skills],
        "ai_profile": _ser(ai_profile),
        "resume_summary": _ser(resume_summary),
        "intelligence_insight": _ser(intelligence),
        "score": _ser(score),
        "outcome": _ser(outcome),
    }


@router.post("/candidates", status_code=201)
def create_candidate(payload: dict, user: dict = Depends(get_current_user)) -> dict:
    """Create a basic candidate record."""
    required = ["first_name", "last_name", "primary_email"]
    for f in required:
        if not payload.get(f):
            raise HTTPException(400, f"{f} is required")

    cid = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    with _db() as cur:
        cur.execute(
            """INSERT INTO candidate
               (id, first_name, middle_name, last_name, primary_email,
                primary_phone_number, status, total_experience, relevant_experience,
                created_on, updated_on, created_by)
               VALUES (%s::uuid, %s, %s, %s, %s, %s, 'Active', %s, %s, %s, %s, %s)""",
            (
                cid,
                payload.get("first_name", ""),
                payload.get("middle_name", ""),
                payload.get("last_name", ""),
                payload.get("primary_email", ""),
                payload.get("phone", ""),
                payload.get("total_experience"),
                payload.get("relevant_experience"),
                now, now,
                user.get("email", "portal"),
            ),
        )
    return {"id": cid, "message": "Candidate created"}


@router.patch("/candidates/{candidate_id}")
def update_candidate(
    candidate_id: str, payload: dict, user: dict = Depends(get_current_user)
) -> dict:
    """Update candidate stage, pipeline_status, assigned_recruiter, interview details."""
    allowed_ai = {"stage", "pipeline_status", "assigned_recruiter", "interview_date",
                  "interview_round", "recruiter_score", "panel_score"}
    ai_updates = {k: v for k, v in payload.items() if k in allowed_ai}

    with _db() as cur:
        if ai_updates:
            set_clause = ", ".join(f"{k}=%s" for k in ai_updates)
            values = list(ai_updates.values())
            cur.execute(
                f"""UPDATE candidate_ai_profile SET {set_clause}, updated_at=NOW()
                    WHERE candidate_id=%s::uuid""",
                values + [candidate_id],
            )
            if cur.rowcount == 0:
                # Create ai_profile row if missing
                cur.execute(
                    """INSERT INTO candidate_ai_profile (candidate_id, stage)
                       VALUES (%s::uuid, 'SOURCED')
                       ON CONFLICT (candidate_id) DO NOTHING""",
                    (candidate_id,),
                )
                if ai_updates:
                    cur.execute(
                        f"""UPDATE candidate_ai_profile SET {set_clause}, updated_at=NOW()
                            WHERE candidate_id=%s::uuid""",
                        values + [candidate_id],
                    )

        # Base candidate updates
        base_allowed = {"status", "primary_phone_number", "address", "city"}
        base_updates = {k: v for k, v in payload.items() if k in base_allowed}
        if base_updates:
            set_clause = ", ".join(f"{k}=%s" for k in base_updates)
            values = list(base_updates.values())
            cur.execute(
                f"UPDATE candidate SET {set_clause}, updated_on=NOW() WHERE id=%s::uuid",
                values + [candidate_id],
            )

    return {"message": "Updated", "updated_fields": list({**ai_updates}.keys())}


@router.post("/candidates/{candidate_id}/analyze")
async def analyze_candidate(
    candidate_id: str, payload: dict = None, user: dict = Depends(get_current_user)
) -> dict:
    """
    Trigger AI analysis for a candidate.
    Reads parsed_resume from candidate_resume_summary OR from payload.
    Updates candidate_ai_profile with scores + analysis.
    """
    if payload is None:
        payload = {}

    # Fetch candidate data
    with _db() as cur:
        cur.execute(
            "SELECT * FROM candidate WHERE id = %s::uuid", (candidate_id,)
        )
        cand = _row(cur)
        if not cand:
            raise HTTPException(404, "Candidate not found")

        cur.execute(
            "SELECT parsed_resume FROM candidate_resume_summary WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        rs = _row(cur)
        parsed_resume = rs["parsed_resume"] if rs and rs.get("parsed_resume") else None

    # Use provided parsed_resume if no stored one
    if not parsed_resume and payload.get("parsed_resume"):
        parsed_resume = payload["parsed_resume"]

    if not parsed_resume:
        raise HTTPException(
            422,
            "No parsed resume available. Upload resume first or provide parsed_resume in body.",
        )

    # Run AI analysis
    try:
        from engine import analyze_resume as _analyze
        from rubric_engine import compute_rubric_score
        from llm_recruiter_analysis import generate_recruiter_analysis

        result = _analyze(parsed_resume)

        # Extract key scores
        scores = result.get("scores", {}) or {}
        rubric = result.get("rubric", {}) or {}

        resume_score = float(scores.get("resume_score") or rubric.get("total_score") or 0)
        band = result.get("band") or rubric.get("band") or ""
        role_family = result.get("role_family") or ""
        dna_fit = result.get("dna") or result.get("dna_fit") or ""
        yoe = float(result.get("yoe") or 0)
        archetype = rubric.get("archetype") or ""
        red_flags = rubric.get("red_flags") or {}
        analysis_json = result

        # Store in candidate_ai_profile
        with _db() as cur:
            cur.execute(
                """INSERT INTO candidate_ai_profile
                   (candidate_id, resume_score, combined_score, band, role_family,
                    dna_fit, yoe, archetype, red_flags_json, analysis_json, updated_at)
                   VALUES (%s::uuid, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                   ON CONFLICT (candidate_id) DO UPDATE SET
                     resume_score   = EXCLUDED.resume_score,
                     combined_score = EXCLUDED.combined_score,
                     band           = EXCLUDED.band,
                     role_family    = EXCLUDED.role_family,
                     dna_fit        = EXCLUDED.dna_fit,
                     yoe            = EXCLUDED.yoe,
                     archetype      = EXCLUDED.archetype,
                     red_flags_json = EXCLUDED.red_flags_json,
                     analysis_json  = EXCLUDED.analysis_json,
                     updated_at     = NOW()""",
                (
                    candidate_id, resume_score, resume_score, band,
                    role_family, dna_fit, yoe, archetype,
                    json.dumps(red_flags), json.dumps(analysis_json),
                ),
            )

            # Update candidate_score (production table)
            cur.execute(
                """INSERT INTO candidate_score (id, candidate_id, score, max_score,
                   strength, drop_flag, created_on, updated_on)
                   VALUES (gen_random_uuid(), %s::uuid, %s, 100, %s, FALSE, NOW(), NOW())
                   ON CONFLICT DO NOTHING""",
                (candidate_id, resume_score, band),
            )

        return {
            "message": "Analysis complete",
            "resume_score": resume_score,
            "band": band,
            "role_family": role_family,
            "dna_fit": dna_fit,
            "yoe": yoe,
        }
    except Exception as exc:
        logger.exception("Analysis failed for %s", candidate_id)
        raise HTTPException(500, f"Analysis failed: {exc}")


@router.get("/candidates/{candidate_id}/analysis")
def get_analysis(candidate_id: str, user: dict = Depends(get_current_user)) -> dict:
    """Return the stored AI analysis for a candidate."""
    with _db() as cur:
        cur.execute(
            "SELECT * FROM candidate_ai_profile WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        profile = _row(cur)
    if not profile:
        raise HTTPException(404, "No analysis found for this candidate")

    def _ser(d):
        return {
            k: str(v) if isinstance(v, uuid.UUID) else
               v.isoformat() if isinstance(v, datetime) else v
            for k, v in d.items()
        }

    return _ser(profile)


@router.post("/candidates/{candidate_id}/outcome")
def record_outcome(
    candidate_id: str, payload: dict, user: dict = Depends(get_current_user)
) -> dict:
    """Record placement / rejection outcome for a candidate."""
    with _db() as cur:
        cur.execute(
            """INSERT INTO candidate_outcome
               (candidate_id, outcome, rejection_stage, placed_company, placed_role,
                placed_date, ctc_offered, feedback_notes, recorded_by, created_at, updated_at)
               VALUES (%s::uuid, %s, %s, %s, %s, %s, %s, %s, %s::uuid, NOW(), NOW())
               ON CONFLICT (candidate_id) DO UPDATE SET
                 outcome         = EXCLUDED.outcome,
                 rejection_stage = EXCLUDED.rejection_stage,
                 placed_company  = EXCLUDED.placed_company,
                 placed_role     = EXCLUDED.placed_role,
                 placed_date     = EXCLUDED.placed_date,
                 ctc_offered     = EXCLUDED.ctc_offered,
                 feedback_notes  = EXCLUDED.feedback_notes,
                 recorded_by     = EXCLUDED.recorded_by,
                 updated_at      = NOW()""",
            (
                candidate_id,
                payload.get("outcome", "IN_PROGRESS"),
                payload.get("rejection_stage"),
                payload.get("placed_company"),
                payload.get("placed_role"),
                payload.get("placed_date"),
                payload.get("ctc_offered"),
                payload.get("feedback_notes"),
                user.get("sub"),
            ),
        )
    return {"message": "Outcome recorded"}


@router.get("/candidates/{candidate_id}/outcome")
def get_outcome(candidate_id: str, user: dict = Depends(get_current_user)) -> dict:
    with _db() as cur:
        cur.execute(
            "SELECT * FROM candidate_outcome WHERE candidate_id = %s::uuid",
            (candidate_id,),
        )
        row = _row(cur)
    if not row:
        return {"outcome": "NOT_RECORDED"}
    return {k: str(v) if isinstance(v, (uuid.UUID,)) else v for k, v in row.items()}


# ── Jobs (JDs) ────────────────────────────────────────────────────────────────

@router.get("/jobs")
def list_jobs(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    status: str = Query(""),
    user: dict = Depends(get_current_user),
) -> dict:
    offset = (page - 1) * page_size
    params: list = []
    where = ""
    if status:
        where = "WHERE jd.status = %s"
        params.append(status)

    with _db() as cur:
        cur.execute(f"SELECT COUNT(*) as total FROM job_description jd {where}", params)
        total = _row(cur)["total"]

        cur.execute(
            f"""SELECT jd.id, jd.code, jd.job_type, jd.job_level, jd.status,
                       jd.experience_min_yrs, jd.experience_max_yrs,
                       jd.required_skills, jd.salary_min, jd.salary_max,
                       jd.total_positions, jd.total_positions_selected,
                       jd.created_on, jd.closed_on, jd.location_city, jd.location_country,
                       jd.department, jd.role_summary,
                       jt.name as job_title, co.name as company_name,
                       u.first_name || ' ' || u.last_name as assigned_to_name
                FROM job_description jd
                LEFT JOIN job_title jt ON jt.id = jd.job_title_id
                LEFT JOIN company co ON co.id = jd.company_id
                LEFT JOIN "user" u ON u.id = jd.assigned_to_user_id
                {where}
                ORDER BY jd.created_on DESC
                LIMIT %s OFFSET %s""",
            params + [page_size, offset],
        )
        rows = _rows(cur)

    def _ser(r):
        return {
            k: str(v) if isinstance(v, uuid.UUID) else
               v.isoformat() if isinstance(v, datetime) else v
            for k, v in r.items()
        }

    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": (total + page_size - 1) // page_size,
        "jobs": [_ser(r) for r in rows],
    }


@router.post("/jobs", status_code=201)
def create_job(payload: dict, user: dict = Depends(get_current_user)) -> dict:
    """Create a job description."""
    jd_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    with _db() as cur:
        cur.execute(
            """INSERT INTO job_description
               (id, code, job_type, job_level, job_description_text,
                required_skills, good_to_have_skills,
                experience_min_yrs, experience_max_yrs,
                total_positions, status, location_city, location_country,
                department, role_summary, created_on, updated_on, created_by)
               VALUES (%s::uuid, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'Open', %s, %s, %s, %s, %s, %s, %s)""",
            (
                jd_id,
                payload.get("code", f"JD-{jd_id[:8].upper()}"),
                payload.get("job_type", "Full Time"),
                payload.get("job_level", ""),
                payload.get("description", ""),
                payload.get("required_skills", []),
                payload.get("good_to_have_skills", []),
                payload.get("yoe_min"),
                payload.get("yoe_max"),
                payload.get("total_positions", 1),
                payload.get("location_city", ""),
                payload.get("location_country", "India"),
                payload.get("department", ""),
                payload.get("role_summary", ""),
                now, now,
                user.get("email", "portal"),
            ),
        )
    return {"id": jd_id, "message": "JD created"}


@router.get("/jobs/{jd_id}")
def get_job(jd_id: str, user: dict = Depends(get_current_user)) -> dict:
    with _db() as cur:
        cur.execute(
            """SELECT jd.*, jt.name as job_title, co.name as company_name
               FROM job_description jd
               LEFT JOIN job_title jt ON jt.id = jd.job_title_id
               LEFT JOIN company co ON co.id = jd.company_id
               WHERE jd.id = %s::uuid""",
            (jd_id,),
        )
        row = _row(cur)
        if not row:
            raise HTTPException(404, "JD not found")

        # Candidate matches for this JD
        cur.execute(
            """SELECT cj.candidate_id, cj.status,
                      c.first_name, c.last_name, c.primary_email,
                      cje.jd_overall_match_score, cje.skill_match_score,
                      ap.resume_score, ap.band, ap.role_family
               FROM candidate_job cj
               JOIN candidate c ON c.id = cj.candidate_id
               LEFT JOIN candidate_job_evaluation cje ON cje.candidate_job_id = cj.id
               LEFT JOIN candidate_ai_profile ap ON ap.candidate_id = cj.candidate_id
               WHERE cj.job_id = %s::uuid
               ORDER BY cje.jd_overall_match_score DESC NULLS LAST""",
            (jd_id,),
        )
        matches = _rows(cur)

    def _ser(d):
        if d is None:
            return None
        return {
            k: str(v) if isinstance(v, uuid.UUID) else
               v.isoformat() if isinstance(v, datetime) else v
            for k, v in d.items()
        }

    return {
        "job": _ser(row),
        "matches": [_ser(m) for m in matches],
    }


@router.patch("/jobs/{jd_id}")
def update_job(
    jd_id: str, payload: dict, user: dict = Depends(get_current_user)
) -> dict:
    allowed = {
        "status", "total_positions", "role_summary", "experience_min_yrs",
        "experience_max_yrs", "location_city", "hiring_deadline",
        "assigned_to_user_id", "department",
    }
    updates = {k: v for k, v in payload.items() if k in allowed}
    if not updates:
        raise HTTPException(400, "No valid fields to update")
    set_clause = ", ".join(f"{k}=%s" for k in updates)
    with _db() as cur:
        cur.execute(
            f"UPDATE job_description SET {set_clause}, updated_on=NOW() WHERE id=%s::uuid",
            list(updates.values()) + [jd_id],
        )
    return {"message": "JD updated"}


@router.delete("/jobs/{jd_id}")
def close_job(jd_id: str, user: dict = Depends(get_current_user)) -> dict:
    with _db() as cur:
        cur.execute(
            "UPDATE job_description SET status='Closed', closed_on=NOW(), updated_on=NOW() WHERE id=%s::uuid",
            (jd_id,),
        )
    return {"message": "JD closed"}


# ── Matching ──────────────────────────────────────────────────────────────────

@router.post("/jobs/{jd_id}/match_all")
async def match_all(jd_id: str, user: dict = Depends(get_current_user)) -> dict:
    """Trigger JD matching against all analyzed candidates."""
    with _db() as cur:
        cur.execute(
            """SELECT jd.*, jt.name as job_title FROM job_description jd
               LEFT JOIN job_title jt ON jt.id = jd.job_title_id
               WHERE jd.id = %s::uuid""",
            (jd_id,),
        )
        jd = _row(cur)
        if not jd:
            raise HTTPException(404, "JD not found")

        # All candidates with AI analysis
        cur.execute(
            """SELECT c.id, ap.analysis_json, ap.role_family, ap.band, ap.resume_score
               FROM candidate c
               JOIN candidate_ai_profile ap ON ap.candidate_id = c.id
               WHERE ap.analysis_json IS NOT NULL""",
        )
        candidates = _rows(cur)

    if not candidates:
        return {"message": "No analyzed candidates found", "matched": 0}

    matched = 0
    try:
        from jd_matching_bridge import match_candidate_to_jd
    except ImportError:
        match_candidate_to_jd = None

    results = []
    for cand in candidates:
        try:
            cid = str(cand["id"])
            analysis = cand["analysis_json"]
            if isinstance(analysis, str):
                analysis = json.loads(analysis)

            jd_match_score = 0.0
            recommendation = "REVIEW"

            if match_candidate_to_jd:
                m = match_candidate_to_jd(analysis, jd)
                jd_match_score = m.get("match_score", 0) or 0
                recommendation = m.get("recommendation", "REVIEW") or "REVIEW"

            combined = (
                0.6 * float(cand.get("resume_score") or 0)
                + 0.4 * jd_match_score
            )

            with _db() as cur:
                # Upsert candidate_job
                cur.execute(
                    """INSERT INTO candidate_job
                       (id, candidate_id, job_id, status, created_on, updated_on, created_by)
                       VALUES (gen_random_uuid(), %s::uuid, %s::uuid, %s, NOW(), NOW(), %s)
                       ON CONFLICT DO NOTHING""",
                    (cid, jd_id, recommendation, user.get("email", "portal")),
                )
                cur.execute(
                    "SELECT id FROM candidate_job WHERE candidate_id=%s::uuid AND job_id=%s::uuid",
                    (cid, jd_id),
                )
                cj = _row(cur)
                if cj:
                    cj_id = cj["id"]
                    cur.execute(
                        """INSERT INTO candidate_job_evaluation
                           (id, candidate_job_id, jd_overall_match_score,
                            skill_match_score, score, max_score, status)
                           VALUES (gen_random_uuid(), %s::uuid, %s, %s, %s, 100, %s)
                           ON CONFLICT DO NOTHING""",
                        (cj_id, jd_match_score, jd_match_score, combined, recommendation),
                    )
            results.append({"candidate_id": cid, "score": combined, "recommendation": recommendation})
            matched += 1
        except Exception as exc:
            logger.warning("Match failed for %s: %s", cand.get("id"), exc)

    return {
        "message": f"Matched {matched} candidates",
        "matched": matched,
        "results": sorted(results, key=lambda x: x["score"], reverse=True)[:20],
    }


@router.get("/candidates/{candidate_id}/match/{jd_id}")
def get_match(
    candidate_id: str, jd_id: str, user: dict = Depends(get_current_user)
) -> dict:
    with _db() as cur:
        cur.execute(
            """SELECT cj.*, cje.*
               FROM candidate_job cj
               LEFT JOIN candidate_job_evaluation cje ON cje.candidate_job_id = cj.id
               WHERE cj.candidate_id = %s::uuid AND cj.job_id = %s::uuid""",
            (candidate_id, jd_id),
        )
        row = _row(cur)
    if not row:
        raise HTTPException(404, "No match found")
    return {k: str(v) if isinstance(v, uuid.UUID) else v for k, v in row.items()}


@router.post("/candidates/{candidate_id}/match/{jd_id}")
async def trigger_match(
    candidate_id: str, jd_id: str, user: dict = Depends(get_current_user)
) -> dict:
    """Trigger JD match for a specific candidate."""
    with _db() as cur:
        cur.execute(
            "SELECT analysis_json, resume_score FROM candidate_ai_profile WHERE candidate_id=%s::uuid",
            (candidate_id,),
        )
        ap = _row(cur)
        if not ap or not ap.get("analysis_json"):
            raise HTTPException(422, "Candidate not analyzed yet. Run /analyze first.")
        cur.execute("SELECT * FROM job_description WHERE id=%s::uuid", (jd_id,))
        jd = _row(cur)
        if not jd:
            raise HTTPException(404, "JD not found")

    analysis = ap["analysis_json"]
    if isinstance(analysis, str):
        analysis = json.loads(analysis)

    jd_match_score = 0.0
    try:
        from jd_matching_bridge import match_candidate_to_jd
        m = match_candidate_to_jd(analysis, jd)
        jd_match_score = m.get("match_score", 0) or 0
        recommendation = m.get("recommendation", "REVIEW") or "REVIEW"
    except Exception:
        recommendation = "REVIEW"

    combined = 0.6 * float(ap.get("resume_score") or 0) + 0.4 * jd_match_score

    with _db() as cur:
        cur.execute(
            """INSERT INTO candidate_job (id, candidate_id, job_id, status, created_on, updated_on, created_by)
               VALUES (gen_random_uuid(), %s::uuid, %s::uuid, %s, NOW(), NOW(), %s)
               ON CONFLICT DO NOTHING""",
            (candidate_id, jd_id, recommendation, user.get("email")),
        )
        cur.execute(
            "SELECT id FROM candidate_job WHERE candidate_id=%s::uuid AND job_id=%s::uuid",
            (candidate_id, jd_id),
        )
        cj = _row(cur)
        if cj:
            cur.execute(
                """INSERT INTO candidate_job_evaluation
                   (id, candidate_job_id, jd_overall_match_score, skill_match_score, score, max_score, status)
                   VALUES (gen_random_uuid(), %s::uuid, %s, %s, %s, 100, %s)
                   ON CONFLICT DO NOTHING""",
                (cj["id"], jd_match_score, jd_match_score, combined, recommendation),
            )

    return {
        "message": "Match complete",
        "jd_match_score": jd_match_score,
        "combined_score": combined,
        "recommendation": recommendation,
    }


# ── Dashboard Stats ───────────────────────────────────────────────────────────

@router.get("/dashboard/stats")
def dashboard_stats(user: dict = Depends(get_current_user)) -> dict:
    with _db() as cur:
        cur.execute("SELECT COUNT(*) as total FROM candidate")
        total_candidates = _row(cur)["total"]

        cur.execute("SELECT COUNT(*) as total FROM candidate_ai_profile WHERE analysis_json IS NOT NULL")
        analyzed = _row(cur)["total"]

        cur.execute("SELECT COUNT(*) as total FROM job_description WHERE status='Open'")
        open_jobs = _row(cur)["total"]

        cur.execute("SELECT COUNT(*) as total FROM candidate_outcome WHERE outcome='PLACED'")
        placed = _row(cur)["total"]

        cur.execute(
            """SELECT stage, COUNT(*) as cnt
               FROM candidate_ai_profile
               GROUP BY stage ORDER BY cnt DESC"""
        )
        stage_breakdown = {r["stage"] or "SOURCED": r["cnt"] for r in _rows(cur)}

        cur.execute(
            """SELECT role_family, COUNT(*) as cnt
               FROM candidate_ai_profile
               WHERE role_family IS NOT NULL AND role_family != ''
               GROUP BY role_family ORDER BY cnt DESC LIMIT 8"""
        )
        role_breakdown = {r["role_family"]: r["cnt"] for r in _rows(cur)}

        cur.execute(
            """SELECT band, COUNT(*) as cnt
               FROM candidate_ai_profile
               WHERE band IS NOT NULL AND band != ''
               GROUP BY band ORDER BY cnt DESC"""
        )
        band_breakdown = {r["band"]: r["cnt"] for r in _rows(cur)}

        cur.execute(
            """SELECT AVG(combined_score) as avg_score,
                      MAX(combined_score) as max_score,
                      MIN(combined_score) as min_score
               FROM candidate_ai_profile
               WHERE combined_score IS NOT NULL"""
        )
        score_stats = _row(cur)

        cur.execute(
            """SELECT jd.id, jt.name as job_title, jd.status,
                      COUNT(cj.id) as candidate_count
               FROM job_description jd
               LEFT JOIN job_title jt ON jt.id = jd.job_title_id
               LEFT JOIN candidate_job cj ON cj.job_id = jd.id
               WHERE jd.status = 'Open'
               GROUP BY jd.id, jt.name, jd.status
               ORDER BY candidate_count DESC LIMIT 5"""
        )
        top_jobs = _rows(cur)

    return {
        "overview": {
            "total_candidates": total_candidates,
            "analyzed_candidates": analyzed,
            "analysis_coverage_pct": round(100 * analyzed / max(total_candidates, 1), 1),
            "open_jobs": open_jobs,
            "placed_candidates": placed,
        },
        "stage_breakdown": stage_breakdown,
        "role_breakdown": role_breakdown,
        "band_breakdown": band_breakdown,
        "score_stats": {
            k: float(v) if v is not None else None
            for k, v in (score_stats or {}).items()
        },
        "top_open_jobs": [
            {k: str(v) if isinstance(v, uuid.UUID) else v for k, v in r.items()}
            for r in top_jobs
        ],
    }


# ── Users ─────────────────────────────────────────────────────────────────────

@router.get("/users")
def list_users(user: dict = Depends(require_admin)) -> list:
    with _db() as cur:
        cur.execute(
            """SELECT id, first_name, last_name, email, role, status,
                      phone_number, department, created_on
               FROM "user" ORDER BY created_on DESC"""
        )
        rows = _rows(cur)
    return [
        {k: str(v) if isinstance(v, uuid.UUID) else
            v.isoformat() if isinstance(v, datetime) else v
         for k, v in r.items()}
        for r in rows
    ]


@router.post("/users/invite", status_code=201)
def invite_user(payload: dict, user: dict = Depends(require_admin)) -> dict:
    """Create a new user with a temporary password."""
    required = ["first_name", "last_name", "email", "role"]
    for f in required:
        if not payload.get(f):
            raise HTTPException(400, f"{f} required")

    email = payload["email"].strip().lower()
    temp_password = payload.get("password", "Tvarah@2025")
    pw_hash = bcrypt.hashpw(temp_password.encode(), bcrypt.gensalt()).decode()

    uid = str(uuid.uuid4())
    with _db() as cur:
        cur.execute('SELECT id FROM "user" WHERE LOWER(email)=%s', (email,))
        if _row(cur):
            raise HTTPException(409, "User already exists")
        cur.execute(
            """INSERT INTO "user"
               (id, first_name, last_name, email, role, status, password_hash, created_on, updated_on)
               VALUES (%s::uuid, %s, %s, %s, %s, 'Active', %s, NOW(), NOW())""",
            (uid, payload["first_name"], payload["last_name"], email, payload["role"], pw_hash),
        )
    return {
        "id": uid,
        "message": "User created",
        "temp_password": temp_password,
    }


@router.patch("/users/{user_id}")
def update_user(
    user_id: str, payload: dict, requester: dict = Depends(require_admin)
) -> dict:
    allowed = {"role", "status", "department", "location", "phone_number"}
    updates = {k: v for k, v in payload.items() if k in allowed}
    if not updates:
        raise HTTPException(400, "No valid fields")
    if payload.get("password"):
        pw = bcrypt.hashpw(payload["password"].encode(), bcrypt.gensalt()).decode()
        updates["password_hash"] = pw
    set_clause = ", ".join(f"{k}=%s" for k in updates)
    with _db() as cur:
        cur.execute(
            f'UPDATE "user" SET {set_clause}, updated_on=NOW() WHERE id=%s::uuid',
            list(updates.values()) + [user_id],
        )
    return {"message": "User updated"}


# ── Reference Data ────────────────────────────────────────────────────────────

@router.get("/ref/skills")
def ref_skills(user: dict = Depends(get_current_user)) -> list:
    with _db() as cur:
        cur.execute(
            "SELECT id, name, category, tier, is_technical FROM skill ORDER BY category, name"
        )
        rows = _rows(cur)
    return [{k: str(v) if isinstance(v, uuid.UUID) else v for k, v in r.items()} for r in rows]


@router.get("/ref/job_titles")
def ref_job_titles(user: dict = Depends(get_current_user)) -> list:
    with _db() as cur:
        cur.execute("SELECT id, name FROM job_title ORDER BY name")
        rows = _rows(cur)
    return [{k: str(v) if isinstance(v, uuid.UUID) else v for k, v in r.items()} for r in rows]


@router.get("/ref/companies")
def ref_companies(
    q: str = Query(""),
    user: dict = Depends(get_current_user),
) -> list:
    with _db() as cur:
        if q:
            cur.execute(
                "SELECT id, name, type, industry FROM company WHERE LOWER(name) LIKE %s ORDER BY name LIMIT 30",
                (f"%{q.lower()}%",),
            )
        else:
            cur.execute("SELECT id, name, type, industry FROM company ORDER BY name LIMIT 100")
        rows = _rows(cur)
    return [{k: str(v) if isinstance(v, uuid.UUID) else v for k, v in r.items()} for r in rows]


# ── Health ────────────────────────────────────────────────────────────────────

@router.get("/health")
def portal_health() -> dict:
    with _db() as cur:
        cur.execute("SELECT 1 as ok")
        _row(cur)
    return {"status": "ok", "service": "tvarah-portal-api", "version": "1.0.0"}
